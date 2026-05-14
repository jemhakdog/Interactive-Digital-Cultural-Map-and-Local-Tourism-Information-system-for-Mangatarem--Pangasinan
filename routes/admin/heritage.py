"""
Admin heritage routes — CRUD for all 5 heritage types.

Admin-only, auto-approved on creation. Uses unified type-based routing
via the heritage registry to avoid code duplication.
"""
import json
import logging
from datetime import date, datetime

from flask import render_template, request, redirect, url_for, flash, abort, jsonify, send_file
from flask_login import login_required, current_user
from extensions import db, limiter
from models import HeritageProfile
from utils.heritage_registry import (
    HERITAGE_TYPES,
    get_heritage_config,
    get_all_types,
    get_display_name,
)
from utils.logger_helper import log_entry, log_success, log_error
from utils.security import detect_sql_injection_attempt, validate_string_input
from . import admin_bp

logger = logging.getLogger(__name__)


def _require_admin():
    """Abort with 403 if current user is not admin."""
    if current_user.role != "admin":
        log_error("admin", "heritage", "Access denied — not admin")
        abort(403)


def _parse_form_value(value, field_type):
    """Parse a form value based on field type with validation."""
    if not value or value.strip() == "":
        return None

    if field_type == "number":
        try:
            return float(value) if "." in value else int(value)
        except (ValueError, TypeError):
            return None
    elif field_type == "date":
        try:
            return datetime.strptime(value, "%Y-%m-%d").date()
        except (ValueError, TypeError):
            return None
    elif field_type == "json":
        try:
            parsed = json.loads(value)
            # Validate JSON size (max 10KB)
            if len(value) > 10 * 1024:
                return None
            return parsed
        except (json.JSONDecodeError, TypeError):
            return value
    else:
        # Validate string fields
        stripped = value.strip()
        # Block SQL injection patterns
        if detect_sql_injection_attempt(stripped):
            return None
        # Enforce max length of 500 for text fields
        is_valid, _ = validate_string_input(stripped, max_length=500)
        if not is_valid:
            return None
        return stripped


def _populate_item_from_form(profile, detail, config, form_data):
    """Populate both profile and detail models from form data."""
    for field_name, label, field_type, required in config["fields"]:
        raw_value = form_data.get(field_name, "")
        parsed = _parse_form_value(raw_value, field_type)
        
        # Determine whether the field belongs to HeritageProfile or the detail model
        if hasattr(HeritageProfile, field_name) and field_name not in ["id"]:
            setattr(profile, field_name, parsed)
        else:
            setattr(detail, field_name, parsed)


def _item_to_dict(profile, detail, config):
    """Convert a heritage item (profile + detail) to a serializable dictionary."""
    result = {
        "id": profile.id,
        "status": profile.status,
        "created_at": profile.created_at.isoformat() if profile.created_at else None,
        "updated_at": profile.updated_at.isoformat() if profile.updated_at else None,
        "user_id": profile.user_id,
        "asset_type": profile.asset_type
    }

    for field_name, label, field_type, required in config["fields"]:
        if hasattr(HeritageProfile, field_name) and field_name not in ["id"]:
            value = getattr(profile, field_name, None)
        else:
            value = getattr(detail, field_name, None)
            
        if isinstance(value, (date, datetime)):
            value = value.isoformat()
        result[field_name] = value

    return result


class ProxyItem:
    """A proxy object to allow templates to access fields seamlessly whether they are on profile or detail."""
    def __init__(self, profile, detail):
        self.profile = profile
        self.detail = detail
        
    def __getattr__(self, name):
        if hasattr(self.detail, name):
            return getattr(self.detail, name)
        elif hasattr(self.profile, name):
            return getattr(self.profile, name)
        raise AttributeError(f"ProxyItem has no attribute '{name}'")


# === Dashboard ===

@admin_bp.route("/heritage")
@login_required
def admin_heritage_dashboard():
    """Overview of all heritage types with submission counts."""
    log_entry("admin", "heritage_dashboard")
    _require_admin()

    type_stats = []
    for slug, config in get_all_types():
        total = HeritageProfile.query.filter_by(asset_type=slug).count()
        approved = HeritageProfile.query.filter_by(asset_type=slug, status="approved").count()
        pending = HeritageProfile.query.filter_by(asset_type=slug, status="pending").count()
        type_stats.append({
            "slug": slug,
            "label": config["label"],
            "label_plural": config["label_plural"],
            "form": config["form"],
            "total": total,
            "approved": approved,
            "pending": pending,
        })

    log_success("admin", "heritage_dashboard", f"Loaded stats for {len(type_stats)} types")
    return render_template(
        "admin/heritage_dashboard.html",
        type_stats=type_stats,
        heritage_types=HERITAGE_TYPES,
    )


# === List ===

@admin_bp.route("/heritage/<heritage_type>")
@login_required
def admin_heritage_list(heritage_type):
    """List all entries for a specific heritage type with pagination."""
    log_entry("admin", "heritage_list", heritage_type=heritage_type)
    _require_admin()

    config = get_heritage_config(heritage_type)
    if not config:
        abort(404)

    model = config["model"]
    page = request.args.get('page', 1, type=int)
    per_page = 20

    # Get all profiles of this type, join with detail, order by profile's created_at
    # We use a joined query to avoid N+1 problem
    paginated = (
        db.session.query(HeritageProfile, model)
        .outerjoin(model, HeritageProfile.id == model.heritage_profile_id)
        .filter(HeritageProfile.asset_type == heritage_type)
        .order_by(HeritageProfile.created_at.desc())
        .paginate(page=page, per_page=per_page, error_out=False)
    )
    
    # Map results to ProxyItem for template compatibility
    items = [ProxyItem(p, d) for p, d in paginated.items if d]

    log_success("admin", "heritage_list", f"Loaded {len(items)} {config['label']} entries (Page {page})")
    return render_template(
        "admin/heritage_list.html",
        items=items,
        pagination=paginated,
        heritage_type=heritage_type,
        config=config,
    )


# === Add (auto-approved) ===

@admin_bp.route("/heritage/<heritage_type>/add", methods=["GET", "POST"])
@login_required
@limiter.limit("10 per minute")
def admin_heritage_add(heritage_type):
    """Create a new heritage entry. Auto-approved since admin creates it."""
    log_entry("admin", "heritage_add", heritage_type=heritage_type, method=request.method)
    _require_admin()

    config = get_heritage_config(heritage_type)
    if not config:
        abort(404)

    if request.method == "POST":
        model = config["model"]
        
        # Create base profile
        profile = HeritageProfile(
            asset_type=heritage_type,
            status="approved",
            user_id=current_user.id
        )
        db.session.add(profile)
        db.session.flush() # To get the profile.id
        
        # Create detail model
        detail = model(heritage_profile_id=profile.id)
        _populate_item_from_form(profile, detail, config, request.form)
        
        db.session.add(detail)
        db.session.commit()

        display_name = get_display_name(detail, heritage_type)
        log_success("admin", "heritage_add", f"Created {config['label']}: '{display_name}'")
        flash(f'{config["label"]} "{display_name}" created successfully!')
        return redirect(url_for("admin.admin_heritage_list", heritage_type=heritage_type))

    return render_template(
        "admin/heritage_form.html",
        heritage_type=heritage_type,
        config=config,
        item=None,
        is_edit=False,
    )


# === Edit ===

@admin_bp.route("/heritage/<heritage_type>/edit/<int:item_id>", methods=["GET", "POST"])
@login_required
@limiter.limit("10 per minute")
def admin_heritage_edit(heritage_type, item_id):
    """Edit an existing heritage entry."""
    log_entry("admin", "heritage_edit", heritage_type=heritage_type, id=item_id, method=request.method)
    _require_admin()

    config = get_heritage_config(heritage_type)
    if not config:
        abort(404)

    model = config["model"]
    profile = HeritageProfile.query.get_or_404(item_id)
    if profile.asset_type != heritage_type:
        abort(404)
        
    detail = model.query.get_or_404(item_id)
    proxy_item = ProxyItem(profile, detail)

    if request.method == "POST":
        _populate_item_from_form(profile, detail, config, request.form)
        db.session.commit()

        display_name = get_display_name(detail, heritage_type)
        log_success("admin", "heritage_edit", f"Updated {config['label']}: '{display_name}'")
        flash(f'{config["label"]} "{display_name}" updated successfully!')
        return redirect(url_for("admin.admin_heritage_list", heritage_type=heritage_type))

    return render_template(
        "admin/heritage_form.html",
        heritage_type=heritage_type,
        config=config,
        item=proxy_item,
        is_edit=True,
    )


# === Delete ===

@admin_bp.route("/heritage/<heritage_type>/delete/<int:item_id>", methods=["POST"])
@login_required
@limiter.limit("10 per minute")
def admin_heritage_delete(heritage_type, item_id):
    """Delete a heritage entry."""
    log_entry("admin", "heritage_delete", heritage_type=heritage_type, id=item_id)
    _require_admin()

    config = get_heritage_config(heritage_type)
    if not config:
        abort(404)

    model = config["model"]
    profile = HeritageProfile.query.get_or_404(item_id)
    detail = model.query.get(item_id)
    
    display_name = get_display_name(detail, heritage_type) if detail else f"Profile {item_id}"

    if detail:
        db.session.delete(detail)
    db.session.delete(profile)
    db.session.commit()

    log_success("admin", "heritage_delete", f"Deleted {config['label']}: '{display_name}'")
    flash(f'{config["label"]} "{display_name}" deleted.')
    return redirect(url_for("admin.admin_heritage_list", heritage_type=heritage_type))


# === API-style JSON endpoints (for AJAX / future frontend) ===

@admin_bp.route("/heritage/<heritage_type>/json")
@login_required
def admin_heritage_json(heritage_type):
    """Return heritage items as JSON with pagination support."""
    _require_admin()

    config = get_heritage_config(heritage_type)
    if not config:
        return jsonify({"error": "Invalid heritage type"}), 404

    model = config["model"]
    page = request.args.get('page', 1, type=int)
    per_page = request.args.get('per_page', 50, type=int)

    paginated = (
        db.session.query(HeritageProfile, model)
        .outerjoin(model, HeritageProfile.id == model.heritage_profile_id)
        .filter(HeritageProfile.asset_type == heritage_type)
        .order_by(HeritageProfile.created_at.desc())
        .paginate(page=page, per_page=per_page, error_out=False)
    )
    
    json_items = [_item_to_dict(p, d, config) for p, d in paginated.items if d]

    return jsonify({
        "heritage_type": heritage_type,
        "label": config["label"],
        "count": len(json_items),
        "total": paginated.total,
        "page": page,
        "pages": paginated.pages,
        "items": json_items,
    })

# === Export to DOCX ===

@admin_bp.route("/heritage/export/docx/<int:profile_id>")
@login_required
def admin_heritage_export_docx(profile_id):
    """Export a heritage record by populating its official template."""
    _require_admin()
    
    from models import HeritageProfile
    from utils.heritage_registry import get_heritage_config
    import os
    from io import BytesIO
    try:
        from docx import Document
    except ImportError:
        flash("Export requires 'python-docx' library.")
        return redirect(url_for('admin.admin_documents'))

    profile = HeritageProfile.query.get_or_404(profile_id)
    config = get_heritage_config(profile.asset_type)
    detail = config["model"].query.get_or_404(profile_id)
    
    # Find template path
    from routes.admin.documents import GATHERED_DIR, FORM_MAPPING
    template_slug = profile.template_slug or profile.asset_type
    meta = FORM_MAPPING.get(template_slug)
    if not meta:
        flash("No template mapping found for this record.")
        return redirect(url_for('admin.admin_documents'))
        
    template_path = os.path.join(GATHERED_DIR, meta["file"])
    if not os.path.exists(template_path):
        flash(f"Template file not found: {meta['file']}")
        return redirect(url_for('admin.admin_documents'))

    try:
        doc = Document(template_path)
        
        # Simple replacement strategy: Iterate through paragraphs and replace labels
        # Heuristic: Find labels in paragraphs and append data
        # We also look for specific tags if any were added
        
        def _get_val(key):
            if hasattr(profile, key) and getattr(profile, key):
                return str(getattr(profile, key))
            if hasattr(detail, key) and getattr(detail, key):
                return str(getattr(detail, key))
            if detail.meta_data and key in detail.meta_data:
                return str(detail.meta_data[key])
            return ""

        # Map labels to field keys
        label_map = {
            "NAME OF NATURAL HERITAGE": "name",
            "NAME OF THE SITE": "name",
            "NAME OF HERITAGE": "name",
            "NAME OF OBJECT": "name",
            "NAME OF THE HERITAGE": "name",
            "B. LOCATION": "location",
            "C. ADDRESS": "address",
            "CONTROL NUMBER": "form_control_number",
            "CATEGORY": "category"
        }

        for para in doc.paragraphs:
            p_text = para.text.upper()
            for label, key in label_map.items():
                if label in p_text and ":" in p_text:
                    val = _get_val(key)
                    if val:
                        # Append value to paragraph if it's empty after colon
                        if para.text.strip().endswith(":"):
                            para.add_run(f" {val}").bold = True
                        elif label in para.text:
                            # Try to find if it already has a value, if not append
                            pass

        # Handle tables (common for field grids)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    c_text = cell.text.upper()
                    for label, key in label_map.items():
                        if label in c_text:
                            val = _get_val(key)
                            if val:
                                # Append to cell or replace if needed
                                cell.text = cell.text + f" {val}"
        
        memory_file = BytesIO()
        doc.save(memory_file)
        memory_file.seek(0)
        
        filename = f"{profile.name_of_asset or 'Export'}_{meta['file']}"
        return send_file(memory_file, download_name=filename, as_attachment=True)
        
    except Exception as e:
        import logging
        logging.error(f"Error generating DOCX: {e}")
        flash(f"Failed to generate export: {str(e)}")
        return redirect(url_for('admin.admin_documents'))
