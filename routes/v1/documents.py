import os
import json
import logging
import zipfile
from datetime import datetime
from io import BytesIO
from flask import Blueprint, render_template, request, send_from_directory, send_file, redirect, url_for, flash
from flask_login import login_required, current_user
from extensions import limiter, db
from utils.logger_helper import log_entry, log_success, log_error
from utils.security import detect_sql_injection_attempt

# Try to import Document from python-docx for the export feature
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

logger = logging.getLogger(__name__)

# Constants for paths
BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
DOCS_DIR = os.path.join(BASE_DIR, "docs", "interview_data")
FORMS_JSON = os.path.join(DOCS_DIR, "forms_structure_analysis.json")
GATHERED_DIR = os.path.join(DOCS_DIR, "gathered_froms")
BACKUP_DIR = os.path.join(DOCS_DIR, "backups")

# Mapping of route slugs to JSON keys and metadata
FORM_MAPPING = {
    "natural": {
        "key": "Form 01A Natural Resources - Land formation 2019 - Mangatarem Tourism",
        "label": "Natural Resources",
        "file": "Form 01A Natural Resources - Land formation 2019 - Mangatarem Tourism.docx",
        "category": "Natural Heritage",
        "heritage_type": "natural"
    },
    "built": {
        "key": "Form 02A Tangible Immovable - Govt and Commercial Buildings 2019 (1) - Mangatarem Tourism",
        "label": "Govt/Commercial Buildings",
        "file": "Form 02A Tangible Immovable - Govt and Commercial Buildings 2019 (1) - Mangatarem Tourism.docx",
        "category": "Built Heritage",
        "heritage_type": "built"
    },
    "movable": {
        "key": "Form 03A Tangible Movable - Archaeological 2019 - Mangatarem Tourism",
        "label": "Archaeological Objects",
        "file": "Form 03A Tangible Movable - Archaeological 2019 - Mangatarem Tourism.docx",
        "category": "Archaeological",
        "heritage_type": "movable"
    },
    "intangible": {
        "key": "Form 04A Intangible Heritage - Oral Tradition and expressions 2019 - Mangatarem Tourism",
        "label": "Oral Traditions",
        "file": "Form 04A Intangible Heritage - Oral Tradition and expressions 2019 - Mangatarem Tourism.docx",
        "category": "Intangible",
        "heritage_type": "intangible"
    },
    "personality": {
        "key": "Form 05 Personalities 2017 - Mangatarem Tourism",
        "label": "Significant Personalities",
        "file": "Form 05 Personalities 2017 - Mangatarem Tourism.docx",
        "category": "Personalities",
        "heritage_type": "personality"
    },
    "institution": {
        "key": "Form 06 Cultural Institutions 2019 - Mangatarem Tourism",
        "label": "Cultural Institutions",
        "file": "Form 06 Cultural Institutions 2019 - Mangatarem Tourism.docx",
        "category": "Institutions",
        "heritage_type": "institution"
    },
    "program": {
        "key": "Form 07 LGU Programs projects for culture 2019 - Mangatarem Tourism",
        "label": "LGU Programs",
        "file": "Form 07 LGU Programs projects for culture 2019 - Mangatarem Tourism.docx",
        "category": "LGU Projects",
        "heritage_type": "program"
    }
}

# Create the v1 documents blueprint
v1_docs_bp = Blueprint("v1_docs", __name__, url_prefix="/admin")


def _require_admin():
    """Abort with 403 if current user is not admin."""
    if current_user.role != "admin":
        log_error("admin", "documents", f"Access denied for user {current_user.username}")
        flash("Access denied.")
        return redirect(url_for("public.index"))
    return None


def _load_all_forms():
    """Load the master forms structure JSON."""
    if not os.path.exists(FORMS_JSON):
        return {}
    try:
        with open(FORMS_JSON, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        logger.error(f"Error loading forms JSON: {e}")
        return {}


def _save_all_forms(data):
    """Save the master forms structure JSON with backup."""
    if not os.path.exists(BACKUP_DIR):
        os.makedirs(BACKUP_DIR, exist_ok=True)
    
    # Create backup of existing JSON
    if os.path.exists(FORMS_JSON):
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        backup_path = os.path.join(BACKUP_DIR, f"forms_structure_{timestamp}.json")
        try:
            with open(FORMS_JSON, 'r', encoding='utf-8') as src:
                with open(backup_path, 'w', encoding='utf-8') as dst:
                    dst.write(src.read())
        except Exception as e:
            logger.warning(f"Could not create backup of forms JSON: {e}")

    try:
        with open(FORMS_JSON, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        return True
    except Exception as e:
        logger.error(f"Error saving forms JSON: {e}")
        return False


def _get_icon_for_cat(category):
    """Return an icon name based on category."""
    cat_lower = category.lower()
    if "natural" in cat_lower: 
        return "folder"
    if "built" in cat_lower: 
        return "image"
    if "archaeological" in cat_lower: 
        return "camera"
    return "folder"


@v1_docs_bp.route("/v1/documents")
@login_required
def v1_documents_view():
    """
    Modernized version of the documents management page (v1).
    Uses real data from the original documents system.
    """
    admin_check = _require_admin()
    if admin_check:
        return admin_check
        
    search_query = request.args.get('q', '').strip().lower()
    logger.info(f"V1 Documents page accessed by user {current_user.username}. Search: {search_query}")
    
    # Load data from original system
    _load_all_forms()
    
    # Fetch recent heritage records
    from models import HeritageProfile
    from sqlalchemy import func
    
    # Get entry counts per asset type
    entry_counts = db.session.query(
        HeritageProfile.asset_type, 
        func.count(HeritageProfile.id)
    ).group_by(HeritageProfile.asset_type).all()
    entry_map = {t: c for t, c in entry_counts if t}
    
    # Fetch recent records for the dashboard
    records = HeritageProfile.query.filter(HeritageProfile.template_slug.isnot(None))\
                                 .order_by(HeritageProfile.created_at.desc())\
                                 .limit(10).all()
    
    # Get all docx files and their stats
    docx_files_stats = []
    if os.path.exists(GATHERED_DIR):
        for f in os.listdir(GATHERED_DIR):
            if f.endswith('.docx'):
                if search_query and search_query not in f.lower():
                    continue
                    
                f_path = os.path.join(GATHERED_DIR, f)
                try:
                    stats = os.stat(f_path)
                    docx_files_stats.append({
                        "name": f,
                        "mtime": stats.st_mtime,
                        "size_kb": round(stats.st_size / 1024, 1),
                        "updated_at": datetime.fromtimestamp(stats.st_mtime).strftime('%b %d, %Y')
                    })
                except OSError:
                    continue
    
    # Sort files by modification time descending
    docx_files_stats.sort(key=lambda x: x["mtime"], reverse=True)
    
    # Process folders (Categories)
    folders_map = {}
    for _, meta in FORM_MAPPING.items():
        cat = meta["category"]
        h_type = meta.get("heritage_type")
        if cat not in folders_map:
            folders_map[cat] = {
                "count": 0, 
                "entries": entry_map.get(h_type, 0) if h_type else 0,
                "icon": _get_icon_for_cat(cat)
            }
        folders_map[cat]["count"] += 1

    # Format folders for template (Folders Details)
    folder_list = []
    for name, info in folders_map.items():
        folder_list.append({
            "name": name,
            "count": info["count"],
            "entries": info["entries"],
            "icon": info["icon"],
            "users": ["BS"]
        })

    # Process documents for the table
    doc_list = []
    for stats in docx_files_stats[:20]:
        label = stats["name"]
        slug = None
        is_shared = False
        for s, meta in FORM_MAPPING.items():
            if meta["file"].split('.')[0] in stats["name"]:
                label = meta["label"]
                slug = s
                is_shared = True
                break
        
        doc_list.append({
            "id": stats["mtime"], 
            "name": stats["name"],
            "label": label,
            "slug": slug,
            "modified": stats["updated_at"],
            "access": "Shared" if is_shared else "Only you",
            "users": ["BS", "JD"] if is_shared else ["BS"]
        })
    
    return render_template(
        "admin/documents_v1.html",
        folders=folder_list,
        documents=doc_list,
        records=records,
        search_query=search_query,
        form_mapping=FORM_MAPPING
    )


@v1_docs_bp.route("/v1/documents/download/<filename>")
@login_required
def v1_document_download(filename):
    """Directly download a docx file."""
    admin_check = _require_admin()
    if admin_check:
        return admin_check
    return send_from_directory(GATHERED_DIR, filename, as_attachment=True)


@v1_docs_bp.route("/v1/documents/category/<path:category_name>")
@login_required
def v1_document_category_files(category_name):
    """List all files associated with a specific form category using modernized v1 layout."""
    admin_check = _require_admin()
    if admin_check:
        return admin_check
    
    # Find all forms in this category
    category_forms = []
    for slug, meta in FORM_MAPPING.items():
        if meta.get("category") == category_name:
            category_forms.append((slug, meta))
            
    if not category_forms:
        flash(f"No forms found for category: {category_name}", "warning")
        return redirect(url_for("v1_docs.v1_documents_view"))

    # Get files matching the pattern in GATHERED_DIR
    files = []
    if os.path.exists(GATHERED_DIR):
        all_files = os.listdir(GATHERED_DIR)
        for slug, meta in category_forms:
            for f in all_files:
                if f.endswith('.docx') and meta["file"].split('.')[0] in f:
                    f_path = os.path.join(GATHERED_DIR, f)
                    try:
                        stats = os.stat(f_path)
                        files.append({
                            "name": f,
                            "updated_at": datetime.fromtimestamp(stats.st_mtime).strftime('%Y-%m-%d %H:%M'),
                            "size": f"{round(stats.st_size / 1024, 1)} KB",
                            "mtime": stats.st_mtime,
                            "slug": slug,
                            "label": meta["label"]
                        })
                    except OSError:
                        continue
    
    # Sort by modification time desc
    files.sort(key=lambda x: x.get("mtime", 0), reverse=True)
    
    return render_template("admin/documents_list_v1.html", 
                           title=f"{category_name} Files", 
                           files=files, 
                           back_url=url_for("v1_docs.v1_documents_view"),
                           is_category_view=True)


@v1_docs_bp.route("/v1/documents/<slug>/files")
@login_required
def v1_document_files(slug):
    """List all files associated with a specific form category using modernized v1 layout."""
    admin_check = _require_admin()
    if admin_check:
        return admin_check
    
    meta = FORM_MAPPING.get(slug)
    if not meta:
        flash("Invalid form type.")
        return redirect(url_for("v1_docs.v1_documents_view"))
    
    # Get files matching the pattern in GATHERED_DIR
    files = []
    if os.path.exists(GATHERED_DIR):
        for f in os.listdir(GATHERED_DIR):
            if f.endswith('.docx') and meta["file"].split('.')[0] in f:
                f_path = os.path.join(GATHERED_DIR, f)
                try:
                    stats = os.stat(f_path)
                    files.append({
                        "name": f,
                        "updated_at": datetime.fromtimestamp(stats.st_mtime).strftime('%Y-%m-%d %H:%M'),
                        "size": f"{round(stats.st_size / 1024, 1)} KB",
                        "slug": slug,
                        "label": meta["label"]
                    })
                except OSError:
                    continue
    
    return render_template("admin/documents_list_v1.html", 
                           title=f"{meta['label']} Files", 
                           files=files, 
                           back_url=url_for("v1_docs.v1_documents_view"))


@v1_docs_bp.route("/v1/documents/<slug>")
@login_required
def v1_document_view(slug):
    """View parsed structure of a specific form using modernized v1 template."""
    admin_check = _require_admin()
    if admin_check:
        return admin_check
    
    meta = FORM_MAPPING.get(slug)
    if not meta:
        flash("Invalid form type.")
        return redirect(url_for("v1_docs.v1_documents_view"))
    
    all_forms_data = _load_all_forms()
    form_data = all_forms_data.get(meta["key"])
    if not form_data:
        flash("Form data not found in structure analysis.")
        return redirect(url_for("v1_docs.v1_documents_view"))
    
    return render_template("admin/documents_view_v1.html", slug=slug, meta=meta, data=form_data)


@v1_docs_bp.route("/v1/documents/<slug>/edit", methods=["GET", "POST"])
@login_required
@limiter.limit("5 per minute")
def v1_document_edit(slug):
    """Edit structure of a specific form using the modernized v1 DOCX-like editor."""
    admin_check = _require_admin()
    if admin_check:
        return admin_check
    
    meta = FORM_MAPPING.get(slug)
    if not meta:
        flash("Invalid form type.")
        return redirect(url_for("v1_docs.v1_documents_view"))
    
    all_forms_data = _load_all_forms()
    form_data = all_forms_data.get(meta["key"])
    if not form_data:
        flash("Form data not found.")
        return redirect(url_for("v1_docs.v1_documents_view"))
    
    if request.method == "POST":
        try:
            # Handle structured data from the new editor
            if request.form.get("structured_data"):
                raw_structured_data = request.form.get("structured_data")
                # Validate size (max 50KB)
                if len(raw_structured_data) > 50 * 1024:
                    flash("Structured data exceeds maximum allowed size (50KB).")
                    return render_template("admin/documents_editor_v1.html",
                                         slug=slug,
                                         meta=meta,
                                         data=form_data)

                new_form_data = json.loads(raw_structured_data)

                # Validate string values in JSON for SQL injection
                def _validate_json_values(obj):
                    """Recursively check all string values in JSON for SQL injection."""
                    if isinstance(obj, str):
                        if detect_sql_injection_attempt(obj):
                            return False
                    elif isinstance(obj, dict):
                        for v in obj.values():
                            if not _validate_json_values(v):
                                return False
                    elif isinstance(obj, list):
                        for item in obj:
                            if not _validate_json_values(item):
                                return False
                    return True

                if not _validate_json_values(new_form_data):
                    flash("Structured data contains invalid patterns.")
                    return render_template("admin/documents_editor_v1.html",
                                         slug=slug,
                                         meta=meta,
                                         data=form_data)

                all_forms_data[meta["key"]] = new_form_data
                if _save_all_forms(all_forms_data):
                    log_success("admin", "admin_document_edit", f"Updated structure for {slug}")
                    flash(f"Document {meta['label']} updated successfully!")
                    return redirect(url_for("v1_docs.v1_document_view", slug=slug))
            else:
                # Fallback for old textarea edit (if still used)
                new_json_str = request.form.get("json_data")
                new_form_data = json.loads(new_json_str)
                all_forms_data[meta["key"]] = new_form_data
                if _save_all_forms(all_forms_data):
                    log_success("admin", "admin_document_edit", f"Updated JSON for {slug}")
                    flash(f"Structure for {meta['label']} updated successfully!")
                    return redirect(url_for("v1_docs.v1_document_view", slug=slug))
        except json.JSONDecodeError:
            flash("Invalid JSON format. Please check your syntax.")
        except Exception as e:
            logger.error(f"Error in document edit: {e}")
            flash(f"An unexpected error occurred: {str(e)}")
            
    return render_template("admin/documents_editor_v1.html", 
                         slug=slug, 
                         meta=meta, 
                         data=form_data)


@v1_docs_bp.route("/v1/documents/<slug>/export")
@login_required
def v1_document_export(slug):
    """Generate and export a DOCX file from the JSON structure."""
    admin_check = _require_admin()
    if admin_check:
        return admin_check
    
    meta = FORM_MAPPING.get(slug)
    if not meta:
        flash("Invalid form type.")
        return redirect(url_for("v1_docs.v1_documents_view"))
    
    all_forms_data = _load_all_forms()
    form_data = all_forms_data.get(meta["key"])
    if not form_data:
        flash("Form data not found for generation.")
        return redirect(url_for("v1_docs.v1_documents_view"))
    
    if not HAS_DOCX:
        flash("Export requires 'python-docx'. Serving original file instead.")
        return send_from_directory(GATHERED_DIR, meta["file"], as_attachment=True)

    try:
        doc = Document()
        # Add paragraphs
        for p in form_data.get("paragraphs", []):
            doc.add_paragraph(p.get("text", ""))
        
        # Add tables
        for t_data in form_data.get("tables", []):
            table = doc.add_table(rows=t_data.get("rows", 0), cols=t_data.get("columns", 0))
            table.style = 'Table Grid'
            content = t_data.get("content", [])
            for r_idx, row_content in enumerate(content):
                for c_idx, cell_text in enumerate(row_content):
                    table.cell(r_idx, c_idx).text = str(cell_text)
                    
        memory_file = BytesIO()
        doc.save(memory_file)
        memory_file.seek(0)
        
        log_success("admin", "admin_document_export", f"Generated DOCX for {slug}")
        return send_file(memory_file, 
                        download_name=f"Generated_{meta['file']}", 
                        as_attachment=True)
                        
    except Exception as e:
        logger.error(f"Error generating DOCX: {e}")
        flash("Failed to generate DOCX. Serving original instead.")
        return send_from_directory(GATHERED_DIR, meta["file"], as_attachment=True)


@v1_docs_bp.route("/v1/documents/export-all")
@login_required
def v1_documents_export_all():
    """Zip all DOCX files and the structure JSON and export."""
    admin_check = _require_admin()
    if admin_check:
        return admin_check
    
    log_entry("admin", "admin_documents_export_all")
    
    memory_file = BytesIO()
    with zipfile.ZipFile(memory_file, 'w', zipfile.ZIP_DEFLATED) as zf:
        # Add DOCX files
        for m in FORM_MAPPING.values():
            f_path = os.path.join(GATHERED_DIR, m["file"])
            if os.path.exists(f_path):
                zf.write(f_path, arcname=f"docx/{m['file']}")
        
        # Add current JSON
        if os.path.exists(FORMS_JSON):
            zf.write(FORMS_JSON, arcname="forms_structure_analysis.json")
            
    memory_file.seek(0)
    
    timestamp = datetime.now().strftime("%Y%m%d")
    return send_file(memory_file, 
                    download_name=f"tourism_forms_backup_{timestamp}.zip", 
                    as_attachment=True)


@v1_docs_bp.route("/v1/documents/import", methods=["POST"])
@login_required
def v1_document_import():
    """Import a DOCX file and update its structure in JSON."""
    admin_check = _require_admin()
    if admin_check:
        return admin_check

    if not HAS_DOCX:
        flash("The 'python-docx' library is not installed. Import feature disabled.")
        return redirect(url_for("v1_docs.v1_documents_view"))

    file = request.files.get("docx_file")
    if not file or not file.filename.endswith(".docx"):
        flash("Please upload a valid .docx file.")
        return redirect(url_for("v1_docs.v1_documents_view"))

    # Validate file size (max 10MB)
    file.seek(0, os.SEEK_END)
    file_size = file.tell()
    file.seek(0)  # Reset pointer
    if file_size > 10 * 1024 * 1024:
        flash("File size exceeds maximum allowed size (10MB).")
        return redirect(url_for("v1_docs.v1_documents_view"))

    flash("Import feature logic is being finalized. Contact support for updates.")
    return redirect(url_for("v1_docs.v1_documents_view"))


@v1_docs_bp.route("/v1/documents/create/<slug>", methods=["GET", "POST"])
@login_required
def v1_document_create(slug):
    """Rapid creation of a heritage record using a modernized document template interface."""
    admin_check = _require_admin()
    if admin_check:
        return admin_check
    
    meta = FORM_MAPPING.get(slug)
    if not meta:
        flash("Invalid document type.")
        return redirect(url_for("v1_docs.v1_documents_view"))
    
    from utils.heritage_registry import get_heritage_config
    heritage_config = get_heritage_config(meta["heritage_type"])
    
    all_forms_data = _load_all_forms()
    template_structure = all_forms_data.get(meta["key"])
    
    if request.method == "POST":
        from models import HeritageProfile
        
        try:
            # Create base profile
            profile = HeritageProfile(
                asset_type=meta["heritage_type"],
                template_slug=slug,
                status="approved", # Admin entries are auto-approved
                user_id=current_user.id
            )
            db.session.add(profile)
            db.session.flush()
            
            # Create detail record
            detail_model = heritage_config["model"]
            detail = detail_model(heritage_profile_id=profile.id)
            
            # Populate fields from form
            extra_data = {}
            for field_name, value in request.form.items():
                if field_name == "csrf_token":
                    continue
                
                # Check if it's a standard field in HeritageProfile or Detail
                if hasattr(profile, field_name):
                    setattr(profile, field_name, value)
                elif hasattr(detail, field_name):
                    setattr(detail, field_name, value)
                else:
                    # Store everything else in meta_data
                    extra_data[field_name] = value
            
            detail.meta_data = extra_data
            db.session.add(detail)
            db.session.commit()
            
            log_success("admin", "document_create", f"Created {slug} record for {profile.name_of_asset}")
            flash("Record created successfully!")
            return redirect(url_for("v1_docs.v1_documents_view"))
            
        except Exception as e:
            db.session.rollback()
            logger.error(f"Error creating document record: {e}")
            flash(f"Error: {str(e)}")
    
    return render_template("admin/documents_rapid_editor_v1.html", 
                           slug=slug, 
                           meta=meta, 
                           template=template_structure,
                           config=heritage_config,
                           is_edit=False)


@v1_docs_bp.route("/v1/documents/record/<int:record_id>/edit", methods=["GET", "POST"])
@login_required
def v1_document_record_edit(record_id):
    """Edit an existing heritage record using the modernized rapid document interface."""
    admin_check = _require_admin()
    if admin_check:
        return admin_check
    
    from models import HeritageProfile
    profile = HeritageProfile.query.get_or_404(record_id)
    slug = profile.template_slug or profile.asset_type # Fallback to asset_type
    
    meta = FORM_MAPPING.get(slug)
    if not meta:
        # Try to find meta by asset_type if slug didn't match
        for s, m in FORM_MAPPING.items():
            if m["heritage_type"] == profile.asset_type:
                meta = m
                slug = s
                break
                
    if not meta:
        flash("Could not find matching template for this record.")
        return redirect(url_for("v1_docs.v1_documents_view"))
    
    from utils.heritage_registry import get_heritage_config
    heritage_config = get_heritage_config(meta["heritage_type"])
    detail = heritage_config["model"].query.get_or_404(record_id)
    
    all_forms_data = _load_all_forms()
    template_structure = all_forms_data.get(meta["key"])
    
    if request.method == "POST":
        try:
            extra_data = detail.meta_data or {}
            for field_name, value in request.form.items():
                if field_name == "csrf_token":
                    continue
                
                if hasattr(profile, field_name):
                    setattr(profile, field_name, value)
                elif hasattr(detail, field_name):
                    setattr(detail, field_name, value)
                else:
                    extra_data[field_name] = value
            
            detail.meta_data = extra_data
            db.session.commit()
            
            log_success("admin", "document_edit", f"Updated {slug} record {record_id}")
            flash("Record updated successfully!")
            return redirect(url_for("v1_docs.v1_documents_view"))
        except Exception as e:
            db.session.rollback()
            flash(f"Error: {str(e)}")
            
    return render_template("admin/documents_rapid_editor_v1.html", 
                           slug=slug, 
                           meta=meta, 
                           template=template_structure,
                           config=heritage_config,
                           record=profile,
                           detail=detail,
                           is_edit=True)
