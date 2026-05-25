import os
from docx import Document

def inspect_docx():
    path = r"d:\porjects\capstone_system\docs\capstone\chapters\full chapters.md.docx"
    if not os.path.exists(path):
        print(f"File not found: {path}")
        return
        
    print(f"Opening DOCX file: {path}")
    print(f"File size: {os.path.getsize(path)} bytes")
    
    try:
        doc = Document(path)
        print(f"Total Paragraphs: {len(doc.paragraphs)}")
        print(f"Total Tables: {len(doc.tables)}")
        
        print("\n--- Structural Outlines (First 20 Non-empty Paragraphs) ---")
        count = 0
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            if text:
                print(f"Para {i} (Style: {p.style.name}): {text[:120]}...")
                count += 1
                if count >= 25:
                    break
                    
    except Exception as e:
        print(f"Error reading file: {e}")

if __name__ == "__main__":
    inspect_docx()
