import docx
import os

def main():
    filepath = r"d:\porjects\capstone_system\docs\interview_data\gathered_froms\Form 07 LGU Programs projects for culture 2019 - Mangatarem Tourism.docx"
    if not os.path.exists(filepath):
        print(f"File not found: {filepath}")
        return

    doc = docx.Document(filepath)
    print("=== PARAGRAPHS ===")
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            print(f"[{i}]: {p.text.strip()}")

    print("\n=== TABLES ===")
    for i, table in enumerate(doc.tables):
        print(f"Table {i}: {len(table.rows)} rows, {len(table.columns)} cols")
        for r_idx, row in enumerate(table.rows):
            row_text = []
            for c_idx, cell in enumerate(row.cells):
                # Avoid duplicate cell references in merged cells
                row_text.append(cell.text.strip().replace("\n", " | "))
            print(f"  Row {r_idx}: {row_text}")

if __name__ == "__main__":
    main()
