import docx
import os

def dump_docx(file_path, output_path):
    if not os.path.exists(file_path):
        print(f"File {file_path} does not exist")
        return
    doc = docx.Document(file_path)
    lines = []
    for p in doc.paragraphs:
        lines.append(p.text)
        
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(row_text))
            
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"Dumped to {output_path}")

dump_docx("docs/rrl/Chapter_1_to_3_Consolidated.docx", "scratch/chapter_1_3.txt")
dump_docx("docs/PROJECT_DOCUMENTATION.docx", "scratch/project_documentation.txt")
