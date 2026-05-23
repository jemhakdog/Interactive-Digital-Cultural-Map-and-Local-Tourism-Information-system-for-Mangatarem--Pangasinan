import sys
import os
import glob
from docx import Document

def main():
    gathered_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "docs", "interview_data", "gathered_froms"))
    docx_files = glob.glob(os.path.join(gathered_dir, "*.docx"))
    
    for filepath in docx_files:
        filename = os.path.basename(filepath)
        print("=" * 80)
        print(f"FILE: {filename}")
        print("=" * 80)
        
        try:
            doc = Document(filepath)
            
            print("FIRST 5 NON-EMPTY PARAGRAPHS:")
            count = 0
            for i, p in enumerate(doc.paragraphs):
                text = p.text.strip()
                if text:
                    print(f"  [{i}]: {text}")
                    count += 1
                    if count >= 5:
                        break
                        
            print("\nTABLE DETAILS:")
            print(f"  Number of tables: {len(doc.tables)}")
            for t_idx, table in enumerate(doc.tables):
                print(f"  Table {t_idx}: {len(table.rows)} rows x {len(table.columns) if table.rows else 0} cols")
                # print first cell text of first few rows
                for r_idx, row in enumerate(table.rows[:3]):
                    cells_text = [c.text.strip().replace('\n', ' ') for c in row.cells]
                    print(f"    Row {r_idx}: {cells_text[:4]}")
                    
        except Exception as e:
            print(f"  Error inspecting file: {e}")
        print("\n")

if __name__ == "__main__":
    main()
