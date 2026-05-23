import docx
import os

def main():
    filepath = r"d:\porjects\capstone_system\docs\interview_data\gathered_froms\Form 07 LGU Programs projects for culture 2019 - Mangatarem Tourism.docx"
    doc = docx.Document(filepath)
    
    print("=== NON-EMPTY PARAGRAPHS AND SURROUNDINGS ===")
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            print(f"[{i}] {text}")
            
    print("\n=== TABLE CELL VALUES ===")
    for t_idx, table in enumerate(doc.tables):
        print(f"\n--- TABLE {t_idx} ---")
        for r_idx, row in enumerate(table.rows):
            row_vals = []
            for c_idx, cell in enumerate(row.cells):
                row_vals.append(cell.text.strip().replace("\n", " [NEWLINE] "))
            print(f"  Row {r_idx}: {row_vals}")

if __name__ == "__main__":
    main()
