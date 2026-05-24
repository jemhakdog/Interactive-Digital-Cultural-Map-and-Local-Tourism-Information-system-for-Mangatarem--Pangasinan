import os
import json
import logging
import zipfile
from datetime import datetime
from io import BytesIO
from flask import Blueprint, render_template, request, send_from_directory, send_file, redirect, url_for, flash
from flask_login import login_required, current_user
from extensions import limiter, db
from utils.logger_helper import log_entry, log_success, log_error
from utils.security import detect_sql_injection_attempt

# Try to import Document from python-docx for the export feature
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

logger = logging.getLogger(__name__)

# Constants for paths
BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
DOCS_DIR = os.path.join(BASE_DIR, "docs", "interview_data")
FORMS_JSON = os.path.join(DOCS_DIR, "forms_structure_analysis.json")
GATHERED_DIR = os.path.join(DOCS_DIR, "gathered_froms")
BACKUP_DIR = os.path.join(DOCS_DIR, "backups")

# Mapping of route slugs to JSON keys and metadata
FORM_MAPPING = {
    "natural": {
        "key": "Form 01A Natural Resources - Land formation 2019 - Mangatarem Tourism",
        "label": "Natural Resources",
        "file": "Form 01A Natural Resources - Land formation 2019 - Mangatarem Tourism.docx",
        "category": "Natural Heritage",
        "heritage_type": "natural"
    },
    "built": {
        "key": "Form 02A Tangible Immovable - Govt and Commercial Buildings 2019 (1) - Mangatarem Tourism",
        "label": "Govt/Commercial Buildings",
        "file": "Form 02A Tangible Immovable - Govt and Commercial Buildings 2019 (1) - Mangatarem Tourism.docx",
        "category": "Built Heritage",
        "heritage_type": "built"
    },
    "movable": {
        "key": "Form 03A Tangible Movable - Archaeological 2019 - Mangatarem Tourism",
        "label": "Archaeological Objects",
        "file": "Form 03A Tangible Movable - Archaeological 2019 - Mangatarem Tourism.docx",
        "category": "Archaeological",
        "heritage_type": "movable"
    },
    "intangible": {
        "key": "Form 04A Intangible Heritage - Oral Tradition and expressions 2019 - Mangatarem Tourism",
        "label": "Oral Traditions",
        "file": "Form 04A Intangible Heritage - Oral Tradition and expressions 2019 - Mangatarem Tourism.docx",
        "category": "Intangible",
        "heritage_type": "intangible"
    },
    "personality": {
        "key": "Form 05 Personalities 2017 - Mangatarem Tourism",
        "label": "Significant Personalities",
        "file": "Form 05 Personalities 2017 - Mangatarem Tourism.docx",
        "category": "Personalities",
        "heritage_type": "personality"
    },
    "institution": {
        "key": "Form 06 Cultural Institutions 2019 - Mangatarem Tourism",
        "label": "Cultural Institutions",
        "file": "Form 06 Cultural Institutions 2019 - Mangatarem Tourism.docx",
        "category": "Institutions",
        "heritage_type": "institution"
    },
    "program": {
        "key": "Form 07 LGU Programs projects for culture 2019 - Mangatarem Tourism",
        "label": "LGU Programs",
        "file": "Form 07 LGU Programs projects for culture 2019 - Mangatarem Tourism.docx",
        "category": "LGU Projects",
        "heritage_type": "program"
    }
}

# Create the v1 documents blueprint
v1_docs_bp = Blueprint("v1_docs", __name__, url_prefix="/admin")


def _parse_docx_file(stream):
    """
    Parses a docx file stream and extracts key-value pairs.
    Returns (heritage_slug, extracted_data).
    """
    try:
        import docx
        from datetime import datetime
        import re
        
        doc = docx.Document(stream)
        
        # Extract paragraphs
        paragraphs_text = []
        for p in doc.paragraphs:
            if p.text.strip():
                paragraphs_text.append(p.text.strip())
                
        # Extract tables
        table_cells = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                table_cells.append(row_text)
                
        # Combine text for auto-detection
        full_text_lower = (" ".join(paragraphs_text) + " " + " ".join([" ".join(row) for row in table_cells])).lower()
        
        # Detect slug based on key phrases matching forms 01-07
        slug = None
        
        # Check first few paragraphs for explicit form prefix to avoid keyword leakage
        first_few_text = " ".join(paragraphs_text[:10]).lower()
        if "form 01" in first_few_text or "form 01a" in first_few_text:
            slug = "natural"
        elif "form 02" in first_few_text or "form 02a" in first_few_text:
            slug = "built"
        elif "form 03" in first_few_text or "form 03a" in first_few_text:
            slug = "movable"
        elif "form 04" in first_few_text or "form 04a" in first_few_text:
            slug = "intangible"
        elif "form 05" in first_few_text:
            slug = "personality"
        elif "form 06" in first_few_text:
            slug = "institution"
        elif "form 07" in first_few_text or "matrix of local government" in first_few_text or "lgu programs" in first_few_text:
            slug = "program"
            
        if not slug:
            # Fallback to reordered global keyword check to avoid leakage (check specific templates first)
            if "lgu programs" in full_text_lower or "local government unit programs" in full_text_lower or "form 07" in full_text_lower or "matrix of local government" in full_text_lower:
                slug = "program"
            elif "cultural institutions" in full_text_lower or "cultural institution" in full_text_lower or "form 06" in full_text_lower:
                slug = "institution"
            elif "significant personalities" in full_text_lower or "personalities" in full_text_lower or "form 05" in full_text_lower:
                slug = "personality"
            elif "oral tradition" in full_text_lower or "intangible cultural" in full_text_lower or "form 04a" in full_text_lower:
                slug = "intangible"
            elif "tangible movable" in full_text_lower or "archaeological" in full_text_lower or "form 03a" in full_text_lower:
                slug = "movable"
            elif "tangible immovable" in full_text_lower or "govt and commercial" in full_text_lower or "form 02a" in full_text_lower:
                slug = "built"
            elif "natural resources" in full_text_lower or "land formation" in full_text_lower or "form 01a" in full_text_lower:
                slug = "natural"
            else:
                slug = "natural" # default fallback
            
        extracted = {}
        
        def clean_val(v):
            return v.strip().strip(":").strip().replace("\xa0", " ")
            
        def match_label(line, labels):
            for label in labels:
                if label.upper() in line.upper():
                    if ":" in line:
                        parts = line.split(":", 1)
                        val = clean_val(parts[1])
                        if val:
                            return val
            return None

        # Parse paragraphs for core values
        for i, line in enumerate(paragraphs_text):
            # Name of asset / site / object / program
            name_val = match_label(line, ["NAME OF NATURAL HERITAGE", "NAME OF THE SITE", "NAME OF HERITAGE", "NAME OF OBJECT", "NAME OF THE HERITAGE", "NAME OF IMMOVABLE HERITAGE", "NAME OF THE ELEMENT", "NAME OF INSTITUTION", "NAME OF PERSONALITY", "NAME", "MUNICIPALITY/CITY", "MUNICIPALITY"])
            if name_val and "name" not in extracted:
                extracted["name"] = name_val
                extracted["name_of_asset"] = name_val
                if "MUNICIPALITY" in line.upper():
                    extracted["program_name"] = name_val + " Cultural Registry Program"
                    extracted["lgu_name"] = name_val
                else:
                    extracted["program_name"] = name_val
                
            # Control number
            ctrl_val = match_label(line, ["CONTROL NUMBER", "CONTROL NO"])
            if ctrl_val and "form_control_number" not in extracted:
                extracted["form_control_number"] = ctrl_val
                
            # Location/Address
            loc_val = match_label(line, ["B. LOCATION", "C. ADDRESS", "C. ADDRESS/LOCATION/COORDINATES", "GEOGRAPHICAL LOCATION", "LOCATION/ADDRESS", "PRESENT ADDRESS", "MUNICIPALITY/CITY"])
            if loc_val:
                if "address" not in extracted: extracted["address"] = loc_val
                if "location" not in extracted: extracted["location"] = loc_val
                if "geographical_range" not in extracted: extracted["geographical_range"] = loc_val
                
            # Birth place
            bp_val = match_label(line, ["BIRTH PLACE"])
            if bp_val and "address" not in extracted:
                extracted["address"] = bp_val
                
            # Prominence/Sub-category
            prom_val = match_label(line, ["PROMINENCE", "TYPE OF CULTURAL INSTITUTION", "CATEGORY", "A. TYPE", "A. SUB-CATEGORY"])
            if prom_val:
                if "category" not in extracted: extracted["category"] = prom_val
                if "prominence_field" not in extracted: extracted["prominence_field"] = prom_val
                if "type_of_natural_heritage" not in extracted: extracted["type_of_natural_heritage"] = prom_val
                if "type_of_object" not in extracted: extracted["type_of_object"] = prom_val
                if "type_of_institution" not in extracted: extracted["type_of_institution"] = prom_val

            # Dates
            dates_val = match_label(line, ["YEAR CONSTRUCTED", "ESTIMATED AGE", "DATE FOUND", "DATE OF BIRTH", "DATE OF DEATH", "DATE CREATED"])
            if dates_val:
                if "dates" not in extracted: extracted["dates"] = dates_val
                if "date_produced" not in extracted: extracted["date_produced"] = dates_val
                if "dates_of_birth_death" not in extracted: extracted["dates_of_birth_death"] = dates_val
                if "date_created" not in extracted: extracted["date_created"] = dates_val
                
            # Ownership
            own_val = match_label(line, ["OWNERSHIP/ JURISDICTION", "OWNERSHIP/JURISDICTION", "OWNERSHIP", "NAME OF OWNER"])
            if own_val and "ownership" not in extracted:
                extracted["ownership"] = own_val
                
            # Look-ahead parser for multi-line descriptive headings
            def look_ahead_text(index, limit=3):
                text_blocks = []
                for j in range(index + 1, min(index + 1 + limit, len(paragraphs_text))):
                    next_line = paragraphs_text[j]
                    
                    # Break if it looks like a new main section header
                    upper_next = next_line.upper().strip()
                    is_new_section = False
                    
                    # Check for roman numerals or alphabet prefixes
                    for prefix in ["I.", "II.", "III.", "IV.", "V.", "VI.", "VII.", "A.", "B.", "C.", "D.", "E.", "F.", "G."]:
                        if upper_next.startswith(prefix):
                            is_new_section = True
                            break
                            
                    # Check for exact header words or standard standalone heading titles
                    for hdr in ["BACKGROUND", "DESCRIPTION", "STORIES", "SIGNIFICANCE", "CONSERVATION MEASURES", "SAFEGUARDING MEASURES", "REFERENCES", "ATTACHMENTS", "LIST OF", "LGU VISION", "LGU MISSION", "LGU GOAL"]:
                        if upper_next == hdr or upper_next.startswith(hdr + " ") or (hdr in upper_next and len(upper_next) < len(hdr) + 5):
                            is_new_section = True
                            break
                            
                    if is_new_section:
                        break
                        
                    if next_line.strip().startswith("(") and next_line.strip().endswith(")"):
                        continue
                    if next_line.strip().startswith("[") and next_line.strip().endswith("]"):
                        continue
                    text_blocks.append(next_line)
                return " ".join(text_blocks).strip()

            if ("II. DESCRIPTION" in line.upper() or "A. PHYSICAL DESCRIPTION" in line.upper() or "A. PHYSICAL FEATURES" in line.upper()) and "description" not in extracted:
                desc = look_ahead_text(i, limit=4)
                if desc: extracted["description"] = desc
            elif ("STORIES ASSOCIATED" in line.upper() or "STORIES/NARRATIVES" in line.upper() or "STORIES AND NARRATIVES" in line.upper()) and "stories" not in extracted:
                stories = look_ahead_text(i, limit=4)
                if stories: extracted["stories"] = stories
            elif ("IV. SIGNIFICANCE" in line.upper() or "BIODIVERSITY SIGNIFICANCE" in line.upper()) and "significance" not in extracted:
                sig = look_ahead_text(i, limit=4)
                if sig: extracted["significance"] = sig
            elif ("CONSTRAINTS/THREATS" in line.upper() or "ISSUES/CHALLENGES" in line.upper()) and "constraints_threats" not in extracted:
                con = look_ahead_text(i, limit=4)
                if con: extracted["constraints_threats"] = con
            elif ("CONSERVATION MEASURES" in line.upper() or "SAFEGUARDING MEASURES" in line.upper()) and "conservation_measures" not in extracted:
                cons_m = look_ahead_text(i, limit=4)
                if cons_m: 
                    extracted["conservation_measures"] = cons_m
                    extracted["safeguarding_description"] = cons_m
            elif "LGU VISION STATEMENT" in line.upper() and "vision" not in extracted:
                vision = look_ahead_text(i, limit=4)
                if vision: extracted["vision"] = vision
            elif "LGU MISSION STATEMENT" in line.upper() and "mission" not in extracted:
                mission = look_ahead_text(i, limit=4)
                if mission: extracted["mission"] = mission
            elif "LGU GOAL STATEMENTS" in line.upper() and "goals" not in extracted:
                goals = look_ahead_text(i, limit=4)
                if goals: extracted["goals"] = goals
            elif "B. BRIEF HISTORY OF THE LGU" in line.upper() and "history" not in extracted:
                history = look_ahead_text(i, limit=4)
                if history:
                    extracted["history"] = history
                    extracted["description"] = history
            elif "G. LGU PROGRAMS ON CULTURE, ARTS, AND HERITAGE" in line.upper() and "strategies" not in extracted:
                strat = look_ahead_text(i, limit=4)
                if strat: extracted["strategies"] = strat

        # Parse tables for details like informants, mapper name, references, profiling date
        for row in table_cells:
            for idx, cell in enumerate(row):
                cell_upper = cell.upper()
                if "KEY INFORMANT/S" in cell_upper or "KEY INFORMANT" in cell_upper:
                    val = ""
                    if ":" in cell:
                        val = clean_val(cell.split(":", 1)[1])
                    if not val and idx + 1 < len(row):
                        val = clean_val(row[idx + 1])
                    if val:
                        extracted["key_informants"] = [val] if slug in ["built", "natural", "movable", "intangible", "personality", "institution", "program"] else val
                        
                elif "REFERENCE/S" in cell_upper or "REFERENCES" in cell_upper or "REFERENCE AND OTHER RESOURCES" in cell_upper:
                    val = ""
                    if ":" in cell:
                        val = clean_val(cell.split(":", 1)[1])
                    if not val and idx + 1 < len(row):
                        val = clean_val(row[idx + 1])
                    if val:
                        extracted["reference_sources"] = val
                        
                elif "NAME OF MAPPER/S" in cell_upper or "NAME OF MAPPER" in cell_upper or "MAPPER" in cell_upper:
                    val = ""
                    if ":" in cell:
                        val = clean_val(cell.split(":", 1)[1])
                    if not val and idx + 1 < len(row):
                        val = clean_val(row[idx + 1])
                    if val:
                        extracted["mapper_name"] = val
                        
                elif "DATE PROFILED" in cell_upper or "DATE OF PROFILE" in cell_upper:
                    val = ""
                    if ":" in cell:
                        val = clean_val(cell.split(":", 1)[1])
                    if not val and idx + 1 < len(row):
                        val = clean_val(row[idx + 1])
                    if val:
                        extracted["date_profiled"] = val
                        
                elif "CONTROL NUMBER" in cell_upper:
                    val = ""
                    if ":" in cell:
                        val = clean_val(cell.split(":", 1)[1])
                    if not val and idx + 1 < len(row):
                        val = clean_val(row[idx + 1])
                    if val:
                        extracted["form_control_number"] = val
                        
                # Check for checkboxes inside table content
                for cell_item in row:
                    if "[x]" in cell_item.lower() or "[ ]" in cell_item:
                        matches = re.findall(r'\[\s*[xX]\s*\]\s*([A-Za-z0-9\s/]+)', cell_item)
                        if matches:
                            checked_val = " ".join([m.strip() for m in matches if m.strip()])
                            if checked_val:
                                if "category" not in extracted or not extracted["category"]:
                                    extracted["category"] = checked_val
                                    extracted["type_of_natural_heritage"] = checked_val
                                    extracted["type_of_object"] = checked_val
                                    extracted["type_of_institution"] = checked_val
                                    
        return slug, extracted
    except Exception as e:
        logger.error(f"Error parsing docx file: {e}")
        import traceback
        traceback.print_exc()
        return None, None


def _require_admin():
    """Abort with 403 if current user is not admin."""
    if current_user.role != "admin":
        log_error("admin", "documents", f"Access denied for user {current_user.username}")
        flash("Access denied.")
        return redirect(url_for("public.index"))
    return None


def _load_all_forms():
    """Load the master forms structure JSON."""
    if not os.path.exists(FORMS_JSON):
        return {}
    try:
        with open(FORMS_JSON, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        logger.error(f"Error loading forms JSON: {e}")
        return {}


def _save_all_forms(data):
    """Save the master forms structure JSON with backup."""
    if not os.path.exists(BACKUP_DIR):
        os.makedirs(BACKUP_DIR, exist_ok=True)
    
    # Create backup of existing JSON
    if os.path.exists(FORMS_JSON):
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        backup_path = os.path.join(BACKUP_DIR, f"forms_structure_{timestamp}.json")
        try:
            with open(FORMS_JSON, 'r', encoding='utf-8') as src:
                with open(backup_path, 'w', encoding='utf-8') as dst:
                    dst.write(src.read())
        except Exception as e:
            logger.warning(f"Could not create backup of forms JSON: {e}")

    try:
        with open(FORMS_JSON, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        return True
    except Exception as e:
        logger.error(f"Error saving forms JSON: {e}")
        return False


def _get_icon_for_cat(category):
    """Return an icon name based on category."""
    cat_lower = category.lower()
    if "natural" in cat_lower: 
        return "folder"
    if "built" in cat_lower: 
        return "image"
    if "archaeological" in cat_lower: 
        return "camera"
    return "folder"


@v1_docs_bp.route("/v1/documents")
@login_required
def v1_documents_view():
    """
    Modernized version of the documents management page (v1).
    Uses real data from the original documents system.
    """
    admin_check = _require_admin()
    if admin_check:
        return admin_check
        
    search_query = request.args.get('q', '').strip().lower()
    logger.info(f"V1 Documents page accessed by user {current_user.username}. Search: {search_query}")
    
    # Load data from original system
    _load_all_forms()
    
    # Fetch recent heritage records
    from models import HeritageProfile
    from sqlalchemy import func
    
    # Get entry counts per asset type
    entry_counts = db.session.query(
        HeritageProfile.asset_type, 
        func.count(HeritageProfile.id)
    ).group_by(HeritageProfile.asset_type).all()
    entry_map = {t: c for t, c in entry_counts if t}
    
    # Fetch recent records for the dashboard
    records = HeritageProfile.query.filter(HeritageProfile.template_slug.isnot(None))\
                                 .order_by(HeritageProfile.created_at.desc())\
                                 .limit(10).all()
    
    # Get all docx files and their stats
    docx_files_stats = []
    if os.path.exists(GATHERED_DIR):
        for f in os.listdir(GATHERED_DIR):
            if f.endswith('.docx'):
                if search_query and search_query not in f.lower():
                    continue
                    
                f_path = os.path.join(GATHERED_DIR, f)
                try:
                    stats = os.stat(f_path)
                    docx_files_stats.append({
                        "name": f,
                        "mtime": stats.st_mtime,
                        "size_kb": round(stats.st_size / 1024, 1),
                        "updated_at": datetime.fromtimestamp(stats.st_mtime).strftime('%b %d, %Y')
                    })
                except OSError:
                    continue
    
    # Sort files by modification time descending
    docx_files_stats.sort(key=lambda x: x["mtime"], reverse=True)
    
    # Process folders (Categories)
    folders_map = {}
    for _, meta in FORM_MAPPING.items():
        cat = meta["category"]
        h_type = meta.get("heritage_type")
        if cat not in folders_map:
            folders_map[cat] = {
                "count": 0, 
                "entries": entry_map.get(h_type, 0) if h_type else 0,
                "icon": _get_icon_for_cat(cat)
            }
        folders_map[cat]["count"] += 1

    # Format folders for template (Folders Details)
    folder_list = []
    for name, info in folders_map.items():
        folder_list.append({
            "name": name,
            "count": info["count"],
            "entries": info["entries"],
            "icon": info["icon"],
            "users": ["BS"]
        })

    # Process documents for the table
    doc_list = []
    for stats in docx_files_stats[:20]:
        label = stats["name"]
        slug = None
        is_shared = False
        for s, meta in FORM_MAPPING.items():
            if meta["file"].split('.')[0] in stats["name"]:
                label = meta["label"]
                slug = s
                is_shared = True
                break
        
        doc_list.append({
            "id": stats["mtime"], 
            "name": stats["name"],
            "label": label,
            "slug": slug,
            "modified": stats["updated_at"],
            "access": "Shared" if is_shared else "Only you",
            "users": ["BS", "JD"] if is_shared else ["BS"]
        })
    
    return render_template(
        "admin/documents_v1.html",
        folders=folder_list,
        documents=doc_list,
        records=records,
        search_query=search_query,
        form_mapping=FORM_MAPPING
    )


@v1_docs_bp.route("/v1/documents/download/<filename>")
@login_required
def v1_document_download(filename):
    """Directly download a docx file."""
    admin_check = _require_admin()
    if admin_check:
        return admin_check
    return send_from_directory(GATHERED_DIR, filename, as_attachment=True)


@v1_docs_bp.route("/v1/documents/category/<path:category_name>")
@login_required
def v1_document_category_files(category_name):
    """List all files associated with a specific form category using modernized v1 layout."""
    admin_check = _require_admin()
    if admin_check:
        return admin_check
    
    # Find all forms in this category
    category_forms = []
    for slug, meta in FORM_MAPPING.items():
        if meta.get("category") == category_name:
            category_forms.append((slug, meta))
            
    if not category_forms:
        flash(f"No forms found for category: {category_name}", "warning")
        return redirect(url_for("v1_docs.v1_documents_view"))

    # Get files matching the pattern in GATHERED_DIR
    files = []
    if os.path.exists(GATHERED_DIR):
        all_files = os.listdir(GATHERED_DIR)
        for slug, meta in category_forms:
            for f in all_files:
                if f.endswith('.docx') and meta["file"].split('.')[0] in f:
                    f_path = os.path.join(GATHERED_DIR, f)
                    try:
                        stats = os.stat(f_path)
                        files.append({
                            "name": f,
                            "updated_at": datetime.fromtimestamp(stats.st_mtime).strftime('%Y-%m-%d %H:%M'),
                            "size": f"{round(stats.st_size / 1024, 1)} KB",
                            "mtime": stats.st_mtime,
                            "slug": slug,
                            "label": meta["label"]
                        })
                    except OSError:
                        continue
    
    # Sort by modification time desc
    files.sort(key=lambda x: x.get("mtime", 0), reverse=True)
    
    return render_template("admin/documents_list_v1.html", 
                           title=f"{category_name} Files", 
                           files=files, 
                           back_url=url_for("v1_docs.v1_documents_view"),
                           is_category_view=True)


@v1_docs_bp.route("/v1/documents/<slug>/files")
@login_required
def v1_document_files(slug):
    """List all files associated with a specific form category using modernized v1 layout."""
    admin_check = _require_admin()
    if admin_check:
        return admin_check
    
    meta = FORM_MAPPING.get(slug)
    if not meta:
        flash("Invalid form type.")
        return redirect(url_for("v1_docs.v1_documents_view"))
    
    # Get files matching the pattern in GATHERED_DIR
    files = []
    if os.path.exists(GATHERED_DIR):
        for f in os.listdir(GATHERED_DIR):
            if f.endswith('.docx') and meta["file"].split('.')[0] in f:
                f_path = os.path.join(GATHERED_DIR, f)
                try:
                    stats = os.stat(f_path)
                    files.append({
                        "name": f,
                        "updated_at": datetime.fromtimestamp(stats.st_mtime).strftime('%Y-%m-%d %H:%M'),
                        "size": f"{round(stats.st_size / 1024, 1)} KB",
                        "slug": slug,
                        "label": meta["label"]
                    })
                except OSError:
                    continue
    
    return render_template("admin/documents_list_v1.html", 
                           title=f"{meta['label']} Files", 
                           files=files, 
                           back_url=url_for("v1_docs.v1_documents_view"))


@v1_docs_bp.route("/v1/documents/<slug>")
@login_required
def v1_document_view(slug):
    """View parsed structure of a specific form using modernized v1 template."""
    admin_check = _require_admin()
    if admin_check:
        return admin_check
    
    meta = FORM_MAPPING.get(slug)
    if not meta:
        flash("Invalid form type.")
        return redirect(url_for("v1_docs.v1_documents_view"))
    
    all_forms_data = _load_all_forms()
    form_data = all_forms_data.get(meta["key"])
    if not form_data:
        flash("Form data not found in structure analysis.")
        return redirect(url_for("v1_docs.v1_documents_view"))
    
    return render_template("admin/documents_view_v1.html", slug=slug, meta=meta, data=form_data)


@v1_docs_bp.route("/v1/documents/<slug>/edit", methods=["GET", "POST"])
@login_required
@limiter.limit("5 per minute")
def v1_document_edit(slug):
    """Edit structure of a specific form using the modernized v1 DOCX-like editor."""
    admin_check = _require_admin()
    if admin_check:
        return admin_check
    
    meta = FORM_MAPPING.get(slug)
    if not meta:
        flash("Invalid form type.")
        return redirect(url_for("v1_docs.v1_documents_view"))
    
    all_forms_data = _load_all_forms()
    form_data = all_forms_data.get(meta["key"])
    if not form_data:
        flash("Form data not found.")
        return redirect(url_for("v1_docs.v1_documents_view"))
    
    if request.method == "POST":
        try:
            # Handle structured data from the new editor
            if request.form.get("structured_data"):
                raw_structured_data = request.form.get("structured_data")
                # Validate size (max 50KB)
                if len(raw_structured_data) > 50 * 1024:
                    flash("Structured data exceeds maximum allowed size (50KB).")
                    return render_template("admin/documents_editor_v1.html",
                                         slug=slug,
                                         meta=meta,
                                         data=form_data)

                new_form_data = json.loads(raw_structured_data)

                # Validate string values in JSON for SQL injection
                def _validate_json_values(obj):
                    """Recursively check all string values in JSON for SQL injection."""
                    if isinstance(obj, str):
                        if detect_sql_injection_attempt(obj):
                            return False
                    elif isinstance(obj, dict):
                        for v in obj.values():
                            if not _validate_json_values(v):
                                return False
                    elif isinstance(obj, list):
                        for item in obj:
                            if not _validate_json_values(item):
                                return False
                    return True

                if not _validate_json_values(new_form_data):
                    flash("Structured data contains invalid patterns.")
                    return render_template("admin/documents_editor_v1.html",
                                         slug=slug,
                                         meta=meta,
                                         data=form_data)

                all_forms_data[meta["key"]] = new_form_data
                if _save_all_forms(all_forms_data):
                    log_success("admin", "admin_document_edit", f"Updated structure for {slug}")
                    flash(f"Document {meta['label']} updated successfully!")
                    return redirect(url_for("v1_docs.v1_document_view", slug=slug))
            else:
                # Fallback for old textarea edit (if still used)
                new_json_str = request.form.get("json_data")
                new_form_data = json.loads(new_json_str)
                all_forms_data[meta["key"]] = new_form_data
                if _save_all_forms(all_forms_data):
                    log_success("admin", "admin_document_edit", f"Updated JSON for {slug}")
                    flash(f"Structure for {meta['label']} updated successfully!")
                    return redirect(url_for("v1_docs.v1_document_view", slug=slug))
        except json.JSONDecodeError:
            flash("Invalid JSON format. Please check your syntax.")
        except Exception as e:
            logger.error(f"Error in document edit: {e}")
            flash(f"An unexpected error occurred: {str(e)}")
            
    return render_template("admin/documents_editor_v1.html", 
                         slug=slug, 
                         meta=meta, 
                         data=form_data)


@v1_docs_bp.route("/v1/documents/<slug>/export")
@login_required
def v1_document_export(slug):
    """Generate and export a DOCX file from the JSON structure."""
    admin_check = _require_admin()
    if admin_check:
        return admin_check
    
    meta = FORM_MAPPING.get(slug)
    if not meta:
        flash("Invalid form type.")
        return redirect(url_for("v1_docs.v1_documents_view"))
    
    all_forms_data = _load_all_forms()
    form_data = all_forms_data.get(meta["key"])
    if not form_data:
        flash("Form data not found for generation.")
        return redirect(url_for("v1_docs.v1_documents_view"))
    
    if not HAS_DOCX:
        flash("Export requires 'python-docx'. Serving original file instead.")
        return send_from_directory(GATHERED_DIR, meta["file"], as_attachment=True)

    try:
        doc = Document()
        # Add paragraphs
        for p in form_data.get("paragraphs", []):
            doc.add_paragraph(p.get("text", ""))
        
        # Add tables
        for t_data in form_data.get("tables", []):
            table = doc.add_table(rows=t_data.get("rows", 0), cols=t_data.get("columns", 0))
            table.style = 'Table Grid'
            content = t_data.get("content", [])
            for r_idx, row_content in enumerate(content):
                for c_idx, cell_text in enumerate(row_content):
                    table.cell(r_idx, c_idx).text = str(cell_text)
                    
        memory_file = BytesIO()
        doc.save(memory_file)
        memory_file.seek(0)
        
        log_success("admin", "admin_document_export", f"Generated DOCX for {slug}")
        return send_file(memory_file, 
                        download_name=f"Generated_{meta['file']}", 
                        as_attachment=True)
                        
    except Exception as e:
        logger.error(f"Error generating DOCX: {e}")
        flash("Failed to generate DOCX. Serving original instead.")
        return send_from_directory(GATHERED_DIR, meta["file"], as_attachment=True)


@v1_docs_bp.route("/v1/documents/export-all")
@login_required
def v1_documents_export_all():
    """Zip all DOCX files and the structure JSON and export."""
    admin_check = _require_admin()
    if admin_check:
        return admin_check
    
    log_entry("admin", "admin_documents_export_all")
    
    memory_file = BytesIO()
    with zipfile.ZipFile(memory_file, 'w', zipfile.ZIP_DEFLATED) as zf:
        # Add DOCX files
        for m in FORM_MAPPING.values():
            f_path = os.path.join(GATHERED_DIR, m["file"])
            if os.path.exists(f_path):
                zf.write(f_path, arcname=f"docx/{m['file']}")
        
        # Add current JSON
        if os.path.exists(FORMS_JSON):
            zf.write(FORMS_JSON, arcname="forms_structure_analysis.json")
            
    memory_file.seek(0)
    
    timestamp = datetime.now().strftime("%Y%m%d")
    return send_file(memory_file, 
                    download_name=f"tourism_forms_backup_{timestamp}.zip", 
                    as_attachment=True)


@v1_docs_bp.route("/v1/documents/import", methods=["POST"])
@login_required
def v1_document_import():
    """Import a DOCX file, auto-detect the form category, parse fields, and pre-fill creator."""
    admin_check = _require_admin()
    if admin_check:
        return admin_check

    if not HAS_DOCX:
        flash("The 'python-docx' library is not installed. Import feature disabled.")
        return redirect(url_for("v1_docs.v1_documents_view"))

    file = request.files.get("docx_file")
    if not file or not file.filename.endswith(".docx"):
        flash("Please upload a valid .docx file.")
        return redirect(url_for("v1_docs.v1_documents_view"))

    # Validate file size (max 10MB)
    file.seek(0, os.SEEK_END)
    file_size = file.tell()
    file.seek(0)  # Reset pointer
    if file_size > 10 * 1024 * 1024:
        flash("File size exceeds maximum allowed size (10MB).")
        return redirect(url_for("v1_docs.v1_documents_view"))

    try:
        from flask import session
        slug, extracted = _parse_docx_file(file.stream)
        if not slug or not extracted:
            flash("Could not parse docx file. Ensure it is a valid Form 01-07 document.", "error")
            return redirect(url_for("v1_docs.v1_documents_view"))

        session["prefilled_heritage_data"] = extracted
        session["prefilled_heritage_slug"] = slug
        
        flash(f"Successfully parsed {FORM_MAPPING.get(slug, {}).get('label', slug)} document! Please review and save.", "success")
        return redirect(url_for("v1_docs.v1_document_create", slug=slug, prefilled=1))
    except Exception as e:
        logger.error(f"Error importing document: {e}")
        flash(f"An error occurred during import: {str(e)}", "error")
        return redirect(url_for("v1_docs.v1_documents_view"))


@v1_docs_bp.route("/v1/documents/create/<slug>", methods=["GET", "POST"])
@login_required
def v1_document_create(slug):
    """Rapid creation of a heritage record using a modernized document template interface."""
    admin_check = _require_admin()
    if admin_check:
        return admin_check
    
    meta = FORM_MAPPING.get(slug)
    if not meta:
        flash("Invalid document type.")
        return redirect(url_for("v1_docs.v1_documents_view"))
    
    from utils.heritage_registry import get_heritage_config
    heritage_config = get_heritage_config(meta["heritage_type"])
    
    all_forms_data = _load_all_forms()
    template_structure = all_forms_data.get(meta["key"])
    
    # Check if we should load pre-filled data from session
    from flask import session
    prefilled_data = None
    if request.method == "GET" and request.args.get("prefilled") == "1" and session.get("prefilled_heritage_slug") == slug:
        prefilled_data = session.pop("prefilled_heritage_data", None)
        session.pop("prefilled_heritage_slug", None)
    
    if request.method == "POST":
        from models import HeritageProfile
        
        try:
            # Create base profile
            profile = HeritageProfile(
                asset_type=meta["heritage_type"],
                template_slug=slug,
                status="approved", # Admin entries are auto-approved
                user_id=current_user.id,
                form_data={}
            )
            
            # Import and use the unified populate logic
            from modules.heritage.admin_routes import _populate_item_from_form
            _populate_item_from_form(profile, heritage_config, request.form)
            
            db.session.add(profile)
            db.session.commit()
            
            log_success("admin", "document_create", f"Created {slug} record for {profile.name_of_asset}")
            flash("Record created successfully!")
            return redirect(url_for("v1_docs.v1_documents_view"))
            
        except Exception as e:
            db.session.rollback()
            logger.error(f"Error creating document record: {e}")
            flash(f"Error: {str(e)}")
    
    return render_template("admin/documents_rapid_editor_v1.html", 
                           slug=slug, 
                           meta=meta, 
                           template=template_structure,
                           config=heritage_config,
                           prefilled_data=prefilled_data,
                           is_edit=False)


@v1_docs_bp.route("/v1/documents/record/<int:record_id>/edit", methods=["GET", "POST"])
@login_required
def v1_document_record_edit(record_id):
    """Edit an existing heritage record using the modernized rapid document interface."""
    admin_check = _require_admin()
    if admin_check:
        return admin_check
    
    from models import HeritageProfile
    profile = HeritageProfile.query.get_or_404(record_id)
    slug = profile.template_slug or profile.asset_type # Fallback to asset_type
    
    meta = FORM_MAPPING.get(slug)
    if not meta:
        # Try to find meta by asset_type if slug didn't match
        for s, m in FORM_MAPPING.items():
            if m["heritage_type"] == profile.asset_type:
                meta = m
                slug = s
                break
                
    if not meta:
        flash("Could not find matching template for this record.")
        return redirect(url_for("v1_docs.v1_documents_view"))
    
    from utils.heritage_registry import get_heritage_config
    heritage_config = get_heritage_config(meta["heritage_type"])
    
    all_forms_data = _load_all_forms()
    template_structure = all_forms_data.get(meta["key"])
    
    if request.method == "POST":
        try:
            from modules.heritage.admin_routes import _populate_item_from_form
            _populate_item_from_form(profile, heritage_config, request.form)
            
            db.session.commit()
            
            log_success("admin", "document_edit", f"Updated {slug} record {record_id}")
            flash("Record updated successfully!")
            return redirect(url_for("v1_docs.v1_documents_view"))
        except Exception as e:
            db.session.rollback()
            flash(f"Error: {str(e)}")
            
    # For compatibility with template, pass ProxyItem
    from modules.heritage.admin_routes import ProxyItem
    proxy_item = ProxyItem(profile)
            
    return render_template("admin/documents_rapid_editor_v1.html", 
                           slug=slug, 
                           meta=meta, 
                           template=template_structure,
                           config=heritage_config,
                           record=proxy_item,
                           detail=proxy_item,
                           is_edit=True)
