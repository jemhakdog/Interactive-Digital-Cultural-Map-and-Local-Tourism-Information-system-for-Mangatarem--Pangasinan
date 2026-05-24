"""
Unit and integration tests for the docx ingestion parser and import routes.
"""

import io
import pytest
from io import BytesIO
from modules.api_v1.documents import _parse_docx_file, FORM_MAPPING

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

@pytest.mark.skipif(not HAS_DOCX, reason="python-docx is not installed")
def test_parse_docx_file_natural_resources():
    """Test parsing a mock Form 01A Natural Resources document."""
    # Create an in-memory docx document
    doc = Document()
    doc.add_paragraph("FORM 01A: NATURAL RESOURCES PROFILE")
    doc.add_paragraph("CONTROL NUMBER: MN-2026-001")
    doc.add_paragraph("NAME OF NATURAL HERITAGE: Mt. Mangatarem")
    doc.add_paragraph("GEOGRAPHICAL LOCATION: Mangatarem, Pangasinan")
    doc.add_paragraph("OWNERSHIP/JURISDICTION: Public Land")
    doc.add_paragraph("II. DESCRIPTION")
    doc.add_paragraph("A beautiful mountain with diverse flora and fauna.")
    doc.add_paragraph("It spans over several barangays.")
    doc.add_paragraph("IV. SIGNIFICANCE")
    doc.add_paragraph("High biodiversity significance.")
    doc.add_paragraph("It is a key watershed area.")
    doc.add_paragraph("CONSERVATION MEASURES")
    doc.add_paragraph("Protected by local forestry laws.")

    # Create a table for informants, mapper, etc.
    table = doc.add_table(rows=4, cols=2)
    table.cell(0, 0).text = "Key Informant:"
    table.cell(0, 1).text = "Juan Dela Cruz"
    table.cell(1, 0).text = "Reference/s:"
    table.cell(1, 1).text = "Mangatarem Tourism Archive"
    table.cell(2, 0).text = "Name of Mapper:"
    table.cell(2, 1).text = "Jane Doe"
    table.cell(3, 0).text = "Date Profiled:"
    table.cell(3, 1).text = "May 22, 2026"

    # Add a row for checkboxes
    table_cb = doc.add_table(rows=1, cols=1)
    table_cb.cell(0, 0).text = "[x] MOUNTAIN / HILL  [ ] CAVE"

    # Save to stream
    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)

    # Run parser
    slug, extracted = _parse_docx_file(stream)

    # Verify category detection and standard mappings
    assert slug == "natural"
    assert extracted["form_control_number"] == "MN-2026-001"
    assert extracted["name"] == "Mt. Mangatarem"
    assert extracted["name_of_asset"] == "Mt. Mangatarem"
    assert extracted["address"] == "Mangatarem, Pangasinan"
    assert extracted["location"] == "Mangatarem, Pangasinan"
    assert extracted["ownership"] == "Public Land"

    # Verify look-ahead multi-line sections
    assert "A beautiful mountain with diverse flora and fauna." in extracted["description"]
    assert "It spans over several barangays." in extracted["description"]
    assert "High biodiversity significance." in extracted["significance"]
    assert "It is a key watershed area." in extracted["significance"]
    assert "Protected by local forestry laws." in extracted["conservation_measures"]

    # Verify grid mapping (informant, reference, mapper, profiling date)
    assert extracted["key_informants"] == ["Juan Dela Cruz"]
    assert extracted["reference_sources"] == "Mangatarem Tourism Archive"
    assert extracted["mapper_name"] == "Jane Doe"
    assert extracted["date_profiled"] == "May 22, 2026"

    # Verify checkbox category extraction
    assert "MOUNTAIN / HILL" in extracted["category"]


@pytest.mark.skipif(not HAS_DOCX, reason="python-docx is not installed")
def test_parse_docx_file_intangible_heritage():
    """Test parsing a mock Form 04A Intangible Heritage document."""
    doc = Document()
    doc.add_paragraph("FORM 04A: INTANGIBLE HERITAGE PROFILE")
    doc.add_paragraph("CONTROL NUMBER: MN-2026-004")
    doc.add_paragraph("NAME OF THE ELEMENT: Mangatarem Native Weaving")
    doc.add_paragraph("II. DESCRIPTION")
    doc.add_paragraph("A traditional weaving method passed down generations.")
    doc.add_paragraph("IV. SIGNIFICANCE")
    doc.add_paragraph("Cultural value showing Mangatarem history.")

    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)

    slug, extracted = _parse_docx_file(stream)
    assert slug == "intangible"
    assert extracted["form_control_number"] == "MN-2026-004"
    assert extracted["name"] == "Mangatarem Native Weaving"
    assert "weaving method" in extracted["description"]
    assert "Cultural value" in extracted["significance"]


@pytest.mark.skipif(not HAS_DOCX, reason="python-docx is not installed")
def test_parse_docx_file_lgu_program():
    """Test parsing a mock Form 07 LGU Programs and Projects for Culture document."""
    doc = Document()
    doc.add_paragraph("FORM 07: MATRIX OF LGU PROGRAMS AND PROJECTS")
    doc.add_paragraph("MUNICIPALITY/CITY: Mangatarem")
    doc.add_paragraph("LGU Vision Statement")
    doc.add_paragraph("To become the leading cultural preservation center in Pangasinan.")
    doc.add_paragraph("LGU Mission Statement")
    doc.add_paragraph("To preserve natural and built heritage of our ancestors.")
    doc.add_paragraph("LGU Goal Statements")
    doc.add_paragraph("Protect 100% of local cultural assets by 2030.")
    doc.add_paragraph("B. BRIEF HISTORY OF THE LGU")
    doc.add_paragraph("Mangatarem was founded in the 19th century.")
    doc.add_paragraph("G. LGU PROGRAMS ON CULTURE, ARTS, AND HERITAGE")
    doc.add_paragraph("Implementation of community-based archiving.")

    table = doc.add_table(rows=4, cols=2)
    table.cell(0, 0).text = "KEY INFORMANT/S:"
    table.cell(0, 1).text = "Mayor Pedro"
    table.cell(1, 0).text = "REFERENCE AND OTHER RESOURCES:"
    table.cell(1, 1).text = "Municipal Executive Order"
    table.cell(2, 0).text = "NAME OF MAPPER/S:"
    table.cell(2, 1).text = "John Mapper"
    table.cell(3, 0).text = "DATE PROFILED:"
    table.cell(3, 1).text = "2026-05-22"

    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)

    slug, extracted = _parse_docx_file(stream)
    assert slug == "program"
    assert extracted["name"] == "Mangatarem"
    assert extracted["lgu_name"] == "Mangatarem"
    assert extracted["program_name"] == "Mangatarem Cultural Registry Program"
    assert extracted["vision"] == "To become the leading cultural preservation center in Pangasinan."
    assert extracted["mission"] == "To preserve natural and built heritage of our ancestors."
    assert extracted["goals"] == "Protect 100% of local cultural assets by 2030."
    assert extracted["history"] == "Mangatarem was founded in the 19th century."
    assert extracted["description"] == "Mangatarem was founded in the 19th century."
    assert extracted["strategies"] == "Implementation of community-based archiving."
    
    # Verify table fields
    assert extracted["key_informants"] == ["Mayor Pedro"]
    assert extracted["reference_sources"] == "Municipal Executive Order"
    assert extracted["mapper_name"] == "John Mapper"
    assert extracted["date_profiled"] == "2026-05-22"

