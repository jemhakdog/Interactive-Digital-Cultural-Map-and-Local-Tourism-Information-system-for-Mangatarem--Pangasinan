from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH


def create_document():
    """Create comprehensive project documentation in DOCX format."""
    doc = Document()

    # Set default font for the document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)

    # Configure Heading styles
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = 'Times New Roman'
        heading_style.font.color.rgb = RGBColor(0, 0, 0)
        heading_style.font.bold = True
        if i == 1:
            heading_style.font.size = Pt(16)
        elif i == 2:
            heading_style.font.size = Pt(14)
        else:
            heading_style.font.size = Pt(12)

    # Helper to add styled paragraphs
    def add_para(text, style='Normal', bold=False, 
                 alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY):
        p = doc.add_paragraph(text, style=style)
        p.alignment = alignment
        if bold:
            for run in p.runs:
                run.bold = True
        return p

    # ==================== TITLE PAGE ====================
    # Add title page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add some spacing at the top
    for _ in range(5):
        doc.add_paragraph()
    
    title_run = title_para.add_run(
        "INTERACTIVE DIGITAL CULTURAL MAP AND LOCAL TOURISM "
        "INFORMATION SYSTEM FOR MANGATAREM, PANGASINAN"
    )
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = 'Times New Roman'
    
    # Add subtitle
    for _ in range(3):
        doc.add_paragraph()
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("A Capstone Project Documentation")
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.name = 'Times New Roman'
    
    # Add team members
    for _ in range(5):
        doc.add_paragraph()
    
    team_para = doc.add_paragraph()
    team_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    team_run = team_para.add_run("Team B3:\n")
    team_run.bold = True
    team_run.font.size = Pt(12)
    
    members = [
        "Jem Carlo Austria",
        "Maryjane Dalas",
        "Rea Solis",
        "Joy De Guzman"
    ]
    for member in members:
        member_para = doc.add_paragraph()
        member_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        member_para.add_run(member)
    
    # Add date
    for _ in range(3):
        doc.add_paragraph()
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run("March 2026")
    
    # Page break after title
    doc.add_page_break()

    # ==================== TABLE OF CONTENTS ====================
    toc_para = doc.add_heading("Table of Contents", level=1)
    toc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("Note: This is a simplified table of contents. "
                     "Page numbers are not auto-generated in this script.")
    
    toc_sections = [
        "Chapter 1: Introduction",
        "  - Background of the Study",
        "  - Purpose and Description",
        "  - Objectives of the Study",
        "  - Conceptual Framework",
        "  - Scope and Limitations",
        "  - Definition of Terms",
        "  - Review of Related Literature",
        "",
        "Chapter 2: Methodology and Design",
        "  - Software Development Methodology",
        "  - Sources of Data",
        "  - Data Gathering Techniques",
        "  - System Architecture",
        "  - System Features and Modules",
        "",
        "Chapter 3: Results and Discussion",
        "  - System Features by User Role",
        "  - Testing and Evaluation Methodologies",
        "",
        "Appendices",
        "  - Administrative Guide",
        "  - User Manual",
        "  - API Reference",
        "  - Deployment Guide",
        "  - Core System Overview"
    ]
    
    for section in toc_sections:
        p = doc.add_paragraph(section)
        if section.startswith("Chapter"):
            p.runs[0].bold = True
    
    doc.add_page_break()

    # ==================== CHAPTER 1 ====================
    doc.add_heading('Chapter 1: Introduction', level=1)

    doc.add_heading('Background of the Study', level=2)
    add_para('The integration of computing solutions in local governance and tourism has become increasingly vital in modernizing public services and promoting cultural heritage. System software and web applications serve as powerful tools to centralize information, streamline processes, and enhance the overall experience for both administrators and end-users. By digitizing cultural data and tourism information, municipalities can ensure wider accessibility, preserve historical records, and promote local attractions more effectively to a broader audience. Embracing such IT infrastructure plans allows organizations to overcome traditional, manual challenges and transition towards a more efficient, interconnected, and dynamic approach to managing local resources.')

    add_para('This study will be conducted for the Local Government Unit (LGU) of Mangatarem, Pangasinan, the main beneficiary and decision-maker for tourism promotion and cultural data management in the municipality. The LGU of Mangatarem plays a central role in driving economic growth through tourism while preserving the rich cultural identity and heritage of the community. As the primary governing body, the LGU is responsible for curating and disseminating accurate information about local landmarks, events, and traditions, ensuring that both residents and visitors have access to reliable resources that reflect the town\'s historical significance.')

    add_para('Currently, the LGU of Mangatarem encounters significant difficulties in managing and promoting tourism information. The existing process is fragmented and largely manual, which results in irregularly updated online content. The lack of standardized tourism materials leads to inconsistent data across different platforms, causing confusion for tourists. Furthermore, slow coordination with stakeholders relying on traditional communication methods delays the sharing of accurate information. This traditional approach also presents limited accessibility for students and researchers who seek reliable cultural and historical information. These challenges establish the need for a centralized platform that can unify and streamline tourism data management.')

    add_para('To address these challenges, the "Interactive Digital Cultural Map and Local Tourism Information System for Mangatarem, Pangasinan" will be developed. This computing solution is introduced as an improved approach to enhancing the organization\'s existing system by replacing fragmented manual processes with a centralized, interactive web-based platform. By digitizing cultural mapping and tourism information, the proposed system aims to provide standardized, easily accessible, and consistently updated data, thereby improving the efficiency of the LGU\'s tourism promotion and enriching the experience of tourists, residents, and researchers alike.')

    doc.add_heading('Purpose and Description', level=2)
    add_para('This Capstone Project was conducted in order to centralize and digitize the tourism and cultural information of Mangatarem, Pangasinan, providing an accessible and interactive platform that streamlines information management and promotes local heritage.')

    add_para('Once the proposed Interactive Digital Cultural Map and Local Tourism Information System for Mangatarem, Pangasinan is implemented to the Local Government Unit (LGU) of Mangatarem, it will hold particular significance for the following beneficiaries:')

    beneficiaries = [
        ("Local Government Unit (LGU) of Mangatarem", 
         "The main beneficiary and decision-maker will benefit from a robust platform for tourism promotion and cultural data management, enabling them to verify and publish accurate information efficiently."),
        ("System Administrators (Tourism Office Staff / IT Staff)", 
         "They will benefit from an administrative dashboard that simplifies the management of user accounts, access permissions, and the approval/rejection of content submissions from contributors."),
        ("Barangay Representatives (Contributors)", 
         "They will benefit from a dedicated portal to submit and update local content, photos, and videos, empowering them to showcase the attractions and events within their respective jurisdictions."),
        ("Public Users (Tourists / Visitors)", 
         "They will benefit from an interactive map that helps them easily locate attractions, search and filter points of interest, and view suggested routes and cultural information for a better travel experience."),
        ("Students and Researchers", 
         "They will benefit from reliable access to historical data, cultural profiles, and community practices, facilitating their academic research and data gathering easily."),
        ("Residents of Mangatarem", 
         "They will gain cultural pride and benefit from the preservation of their heritage through a secure, digital platform that documents their traditions.")
    ]
    
    for i, (name, desc) in enumerate(beneficiaries, 1):
        p = doc.add_paragraph(style='List Number')
        p.add_run(name).bold = True
        p.add_run(f" – {desc}")

    add_para('The rationale of the project is to resolve the inconsistencies, slow communication, and limited accessibility prevalent in the current manual tourism management processes. By standardizing information and leveraging digital mapping technology, the project creates a unified resource for all stakeholders. It is assumed that the proposed computing solution will effectively address the existing problems by providing real-time, accurate updates, fostering better coordination among barangay representatives and the LGU, and offering an engaging, user-friendly interface for public exploration.')

    doc.add_heading('Objectives of the Study', level=2)
    add_para('The main objective of the study is to design and develop an Interactive Digital Cultural Map and Local Tourism Information System for the Local Government Unit (LGU) of Mangatarem, Pangasinan.')

    add_para('Furthermore, the developers aim to achieve the following specific objectives:')

    objectives = [
        "To analyze the existing process of managing and disseminating tourism and cultural information in the municipality to identify inefficiencies, challenges, and opportunities for improvement in information centralization.",
        "To identify the features of the system for the following users:",
        "To test and evaluate the system's functionality, performance, security, usability, and acceptability to ensure it meets user requirements and standards.",
        "To prepare an implementation plan for the deployment of the system."
    ]
    
    for obj in objectives:
        if "users" in obj:
            p = doc.add_paragraph(obj, style='List Number')
            doc.add_paragraph('System Administrator (Tourism/IT Staff)', style='List Bullet')
            doc.add_paragraph('Barangay Representative (Contributor)', style='List Bullet')
            doc.add_paragraph('Public User (Tourists / Visitors)', style='List Bullet')
            doc.add_paragraph('Students and Researchers', style='List Bullet')
        else:
            doc.add_paragraph(obj, style='List Number')

    doc.add_heading('Conceptual Framework', level=2)
    add_para("The Input-Process-Output (IPO) model is utilized to provide a clear and structured representation of the system's development lifecycle. The Input phase defines the foundational prerequisites, encompassing the knowledge, hardware, and software requirements necessary to build the system. The Process phase outlines the systematic Software Development Methodology chosen to transform these inputs into a functional product, detailing the specific stages of development. The Output phase represents the final deliverable, which is the operational computing solution that addresses the needs identified during the analysis. Feedback mechanisms continuously refine the inputs and processes to ensure the output meets the desired standards.")

    add_para('Input includes knowledge requirements (technical skills in web system development using HTML, CSS, JavaScript, PHP, MySQL, and an understanding of tourism mapping and user roles such as admins, contributors, and tourists), hardware requirements (components needed for both development and deployment, such as an Intel Core i5/AMD Ryzen 5, 8GB-16GB RAM, SSD storage, and standard peripherals), and software requirements (software tools such as Visual Studio Code, MySQL, XAMPP/LAMP server, and UI design tools like Figma).')

    add_para('Process refers to the Rapid Application Development (RAD) Methodology to be used, involving Requirements Planning, User Design, Construction, and Cutover.')

    add_para('Output is the expected computing solution: the Interactive Digital Cultural Map and Local Tourism Information System for Mangatarem, Pangasinan.')

    add_para('[Conceptual Framework Diagram - See original Mermaid code for visual representation]', 
             style='Normal', alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)

    add_para('The conceptual framework illustrated above delineates the systematic flow of the project. The Inputs specify the technical and physical resources, along with the domain knowledge required by the developers. These resources feed into the Process, which employs the Rapid Application Development (RAD) methodology to iteratively design, prototype, and build the platform. This structured process guarantees that the final Output—the Interactive Digital Cultural Map—is developed efficiently and aligns with the requirements of the Mangatarem LGU and its stakeholders. The feedback loop ensures that any issues identified post-deployment can be addressed to refine and maintain the system.')

    doc.add_heading('Scope and Limitations', level=2)

    doc.add_heading('Scope', level=3)
    add_para('The project focuses on the development of a web-based Interactive Digital Cultural Map and Local Tourism Information System for the LGU of Mangatarem. The system will be built utilizing web technologies including HTML, CSS, JavaScript, PHP, and MySQL as the primary database management system, with Figma utilized for interface design. The key functionalities provided by the software include a public interactive map for tourists to locate attractions, filter points of interest, and view cultural information. It features a decentralized content contribution portal allowing authorized Barangay Representatives to upload photos, update local history, and add events. Furthermore, the system includes a centralized Admin Dashboard for LGU/Tourism staff to moderate content submissions (approve/reject), manage user accounts, and oversee platform operations. Students and researchers are provided structured access to historical data and cultural profiles to support academic data gathering.')

    doc.add_heading('Limitations', level=3)
    add_para('While the system aims to comprehensive tourism management, it will not include an online booking or payment gateway for local accommodations or tour guides. The system is highly dependent on internet connectivity; thus, full offline capabilities for the interactive map are restricted. Performance and scalability are designed to accommodate the current reasonable volume of tourist traffic and barangay contributions, but extreme surges beyond typical municipal capacity may require future server upgrades. Access to the content contribution and moderation modules is strictly restricted to authorized LGU personnel and registered Barangay Representatives, meaning the general public cannot directly alter map data without LGU approval.')

    doc.add_heading('Definition of Terms', level=2)

    terms = [
        ("Admin/System Administrator:", "Refers to the LGU Tourism Office Staff or IT personnel responsible for reviewing content, managing users, and overseeing the technical maintenance of the system."),
        ("Barangay Representative:", "An authorized contributor who submits, updates, and uploads local tourism and cultural information to the system on behalf of their specific jurisdiction."),
        ("Contributor:", "A user role (typically a Barangay Representative) with permissions to propose new content such as photos, history, and events to the platform."),
        ("Interactive Digital Cultural Map:", "The core feature of the system that provides a visual, geographical representation of tourist spots, landmarks, and cultural heritage sites within Mangatarem."),
        ("Local Government Unit (LGU):", "In this study, it refers to the municipal government of Mangatarem, Pangasinan, serving as the main beneficiary and authoritative body over the tourism system."),
        ("Public User:", "Refers to tourists, visitors, or general users who navigate the interactive map and view the published cultural information without the need for administrative access."),
        ("Rapid Application Development (RAD):", "The selected software development methodology characterized by rapid prototyping, iterative feedback, and flexible requirements gathering to speed up system construction.")
    ]

    for term, definition in terms:
        p = doc.add_paragraph()
        p.add_run(term).bold = True
        p.add_run(f" {definition}")

    doc.add_heading('Review of Related Literature', level=2)

    add_para('(Note: This section requires sourcing 5 local and 5 foreign literature from 2020-2025 aligned with the objectives. Temporary placeholders are provided here to adhere to the formatting structure until formal literature review research is conducted and injected).', bold=True)

    doc.add_heading('Existing Tourism and Cultural Information Management Processes', level=3)
    add_para('Smith (2022) highlights the inefficiencies in traditional municipal tourism management, stating that reliance on fragmented physical records and isolated social media announcements significantly limits the reach and accuracy of cultural promotion. A critical evaluation of Smith\'s work underscores the necessity for municipalities to shift from manual archiving to centralized digital repositories to ensure consistent data availability.')
    add_para('Dela Cruz et al. (2021) observed similar challenges in local Philippine contexts, where LGUs struggle with inconsistent tourism data across various barangays due to the lack of a unified reporting system. The study suggests that empowering local grassroots (barangays) with direct contribution access to a centralized pool can drastically reduce data dissemination delays.')

    doc.add_heading('Key Features of Web-Based Tourism Systems', level=3)
    add_para('Johnson and White (2023) examined the impact of interactive mapping on tourist engagement, finding that digital maps equipped with filtering capabilities and rich multimedia pop-ups increase visitor retention and exploration confidence by 40%. The research emphasizes that an intuitive UI is a critical feature for public-facing tourism platforms to effectively guide tourists.')
    add_para('Reyes (2024) evaluated the role of decentralized content management in e-governance systems. The study determined that creating specific user roles, such as local contributors and central moderators, improves content accuracy and accountability. This points directly to the necessity of the proposed Admin and Barangay Representative roles to ensure the integrity of the published cultural data.')

    doc.add_heading('Usability and Acceptability of Information Systems', level=3)
    add_para('Anderson (2021) provides a comprehensive review of usability testing methodologies for public sector web applications. The author emphasizes that utilizing standardized Likert-scale surveys to measure navigation ease and interface clarity is essential for systems targeting diverse demographics, such as both tech-savvy students and general tourists.')
    add_para('Bautista (2022) explores user acceptance testing (UAT) frameworks in Philippine LGUs adopting new IT infrastructure. The findings suggest that early and iterative involvement of stakeholders (e.g., tourism officers) during the prototyping phase drastically improves the final acceptance score of the system, supporting the use of the RAD methodology for this project.')

    doc.add_heading('System Implementation and Deployment Strategies', level=3)
    add_para('Chen (2023) discusses deployment strategies for cloud-based municipal systems, identifying that a phased rollout paired with comprehensive user training for administrative staff significantly reduces early-stage operational friction. The review points out that a clear implementation timeline and resource allocation plan are crucial for minimizing downtime during the cutover phase.')
    add_para('Gomez et al. (2025) analyzed recent case studies of digital tourism map deployments in rural areas. They highlight that ensuring reliable web hosting and defining clear data governance policies prior to the launch are critical steps. This supports the objective of preparing a detailed and robust implementation plan tailored specifically to the technological readiness of the target LGU.')

    # ==================== CHAPTER 2 ====================
    doc.add_page_break()
    doc.add_heading('Chapter 2: Methodology and Design', level=1)

    add_para('This chapter discusses the methodology and design processes employed in developing the Interactive Digital Cultural Map and Local Tourism Information System for Mangatarem. It details the chosen software development methodology, the sources of data, the data gathering techniques applied, and presents the structural design of the system through architectural diagrams and flowcharts, culminating in a planned implementation strategy.')

    doc.add_heading('Software Development Methodology', level=2)
    add_para('The implementation of a robust Software Development Methodology (SDM) is crucial in software engineering as it provides a structured framework for planning, creating, testing, and deploying an information system. For the development of the Interactive Digital Cultural Map and Local Tourism Information System, a well-defined SDM ensures that the specific requirements of the Mangatarem LGU and its stakeholders are met systematically, minimizing risks and ensuring the timely delivery of a functional product.')

    add_para('The methodology chosen for this study is Rapid Application Development (RAD). This approach was selected because it emphasizes rapid prototyping and iterative delivery over strict planning. Given the dynamic nature of tourism information and the need to accommodate the evolving input from various Barangay Representatives and the LGU, RAD allows the developers to quickly adapt to feedback without disrupting the overall project timeline.')

    add_para('Rapid Application Development is an agile-based methodology characterized by its flexible, user-centric approach. Its key principles involve continuous stakeholder engagement, where users actively participate in reviewing prototypes. This methodology is typically utilized in projects where user interface constraints are critical, and requirements might shift as the stakeholders visualize the developing software. The primary advantage of RAD is its capability to significantly reduce development time while maintaining high user satisfaction through iterative refinement.')

    phases = [
        ("Requirements Planning:", "In this initial phase, the developers collaborated with the LGU staff and potential users to identify the core problems of the fragmented manual system. Key objectives, such as centralizing data and establishing user roles (Admin, Barangay Representative, Public User), were defined."),
        ("User Design (Prototyping):", "Developers created interactive prototypes and wireframes of the Interactive Map, Admin Dashboard, and Contributor Portals using Figma. These designs were presented to stakeholders for iterative feedback, ensuring the UI/UX aligned with tourist and administrative expectations."),
        ("Construction:", "Following the approval of prototypes, the actual coding commenced. Utilizing HTML, CSS, JavaScript for the frontend and PHP/MySQL for the backend, developers built the functional modules in iterative cycles, continuously testing features against the refined requirements."),
        ("Cutover (Testing, Deployment, & Maintenance):", "The final phase involves comprehensive functional, usability, and security testing. Once the system meets the quality standards, it will be deployed to the LGU's live environment, followed by user training for the Tourism Officers and Barangay Representatives, and ongoing technical maintenance.")
    ]

    for i, (title, desc) in enumerate(phases, 1):
        p = doc.add_paragraph(style='List Number')
        p.add_run(title).bold = True
        p.add_run(f" {desc}")

    doc.add_heading('Sources of Data', level=2)
    add_para('The primary sources of data for this project are individuals, groups, and locations within the municipality of Mangatarem that hold crucial tourism and cultural information. Key sources include the LGU Tourism Office Staff, who provide the official policies, existing manual records, and municipal-level tourism initiatives. Barangay Officials and Representatives serve as vital sources for localized cultural data, specific landmarks, and grassroots community events. Additionally, existing municipal archives, physical maps, and local historical documents act as secondary data sources to establish the initial database of the system.')

    doc.add_heading('Data Gathering Techniques', level=2)
    add_para('To ensure comprehensive and accurate requirement analysis, the following data gathering techniques were applied:')

    techniques = [
        ("Interviews:", "Semi-structured interviews were conducted with the Tourism Office Staff and select Barangay Representatives. This technique was chosen to understand HOW the current manual sharing of tourism data operates and WHY delays occur. It was applied during the Requirements Planning phase to gather qualitative insights directly from the administrators."),
        ("Observation:", "The developers observed the daily workflow of the Tourism Office when handling inquiries and content updates. This was done to identify the specific bottlenecks in their traditional communication methods and to determine the necessary features for the Admin Dashboard."),
        ("Document Analysis:", "The team analyzed existing physical tourism brochures, fragmented social media posts, and municipal records. This technique provided a baseline understanding of the current content structure and helped identify gaps in the existing information dissemination process.")
    ]
    
    for title, desc in techniques:
        p = doc.add_paragraph(style='List Number')
        p.add_run(title).bold = True
        p.add_run(f" {desc}")

    doc.add_heading('System Architecture', level=2)
    add_para('The Interactive Digital Cultural Map and Local Tourism Information System for Mangatarem, Pangasinan is built on a modern web-based architecture designed for scalability, performance, and ease of maintenance.')

    doc.add_heading('Tech Stack', level=3)
    add_para('Backend:')
    backend_items = [
        "Python with Flask framework for the core application logic",
        "SQLAlchemy ORM for database interactions",
        "Supabase (PostgreSQL) for cloud database hosting",
        "SQLite for local development environment"
    ]
    for item in backend_items:
        doc.add_paragraph(item, style='List Bullet')
    
    add_para('Frontend:')
    frontend_items = [
        "HTML5, CSS3, and JavaScript for core web technologies",
        "Tailwind CSS for responsive utility-first styling",
        "Leaflet.js for interactive map rendering",
        "Jinja2 templating engine for dynamic content"
    ]
    for item in frontend_items:
        doc.add_paragraph(item, style='List Bullet')
    
    add_para('Deployment:')
    deployment_items = [
        "Vercel for serverless frontend hosting",
        "Supabase for managed PostgreSQL database",
        "Git/GitHub for version control"
    ]
    for item in deployment_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Architectural Pattern', level=3)
    add_para('The system follows the Model-View-Controller (MVC) architectural pattern:')
    mvc_items = [
        ("Model", "Database models defined in models.py using SQLAlchemy, representing entities like User, Attraction, Event, HeritageProfile, etc."),
        ("View", "HTML templates with Jinja2 that render the user interface and display data to users"),
        ("Controller", "Flask route handlers in the routes/ directory that process requests, interact with models, and render views")
    ]
    for name, desc in mvc_items:
        p = doc.add_paragraph(style='List Number')
        p.add_run(name).bold = True
        p.add_run(f" – {desc}")

    doc.add_heading('System Features and Modules', level=2)
    add_para('The Interactive Digital Cultural Map and Local Tourism Information System for Mangatarem, Pangasinan is designed to address the challenges of fragmented manual tourism processes. The system is divided into specific modules tailored to the needs and access privileges of its key stakeholders.')

    doc.add_heading('Public User Module', level=3)
    public_features = [
        "Interactive Map: Explore attractions with markers and popups",
        "Search and Filter: Find locations by category, name, or features",
        "Attraction Details: View photos, descriptions, and location information",
        "Events Calendar: Browse upcoming cultural events and festivals",
        "Gallery: View multimedia content of Mangatarem's heritage",
        "Suggested Routes: Get recommended itineraries for visits"
    ]
    for feature in public_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('Barangay Contributor Module', level=3)
    contributor_features = [
        "Content Submission: Add new attractions, events, and heritage data",
        "Media Upload: Attach photos and videos to submissions",
        "Submission Tracking: Monitor approval status of submitted content",
        "Edit Pending Items: Modify submissions before approval",
        "Barangay Dashboard: Overview of contributed content"
    ]
    for feature in contributor_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('Administrator Module', level=3)
    admin_features = [
        "Content Moderation: Approve or reject contributor submissions",
        "User Management: Create, update, and manage user accounts",
        "Analytics Dashboard: View system metrics and content statistics",
        "Heritage Data Management: Oversee the 5 heritage categories",
        "Form Management: Import and export cultural data forms",
        "System Configuration: Manage categories, tags, and settings"
    ]
    for feature in admin_features:
        doc.add_paragraph(feature, style='List Bullet')

    # ==================== CHAPTER 3 ====================
    doc.add_page_break()
    doc.add_heading('Chapter 3: Results and Discussion', level=1)

    add_para('(Note: As per the BCC BSIT Capstone Project Guide Revised 2025, Capstone 1 focuses exclusively on Chapters 1, 2, and the system design/features of Chapter 3. The actual implementation results, deployment findings, and final evaluation scores will be completed during Capstone 2. This chapter currently outlines the expected system features per user role and the planned testing and evaluation methodologies.)')

    doc.add_heading('System Features by User Role', level=2)

    doc.add_heading('Tourist/Visitor Experience', level=3)
    add_para('The public-facing interface prioritizes ease of navigation and rich visual content. Tourists can access the interactive map immediately without registration. The map displays all verified attractions with color-coded markers based on category (Natural, Cultural, Historical, etc.). Clicking a marker opens a popup with a thumbnail, name, and brief description. Users can filter attractions by type, search by keyword, and view detailed pages with full photo galleries, historical context, and location coordinates.')

    doc.add_heading('Barangay Representative Workflow', level=3)
    add_para('Contributors access a dedicated dashboard after login. The submission process is structured around the five heritage categories: Natural Resources, Built Heritage, Movable Heritage, Intangible Heritage, and Living Heritage. Each category has a tailored form based on the official municipal forms (01A-07). Contributors can upload multiple photos, specify GPS coordinates, and provide detailed descriptions. Submissions enter a pending state until reviewed by the admin.')

    doc.add_heading('Administrator Capabilities', level=3)
    add_para('Administrators have access to a comprehensive dashboard showing pending submissions, total attractions, events, and user counts. The content moderation interface displays submissions with side-by-side comparison of proposed changes. Admins can approve (publishing content immediately) or reject (with comments for the contributor). User management allows creating accounts, resetting passwords, and assigning roles. The heritage management module provides CRUD operations for all five heritage tables.')

    doc.add_heading('Testing and Evaluation Methodologies', level=2)

    doc.add_heading('Functional Testing', level=3)
    add_para('The system will undergo comprehensive functional testing to ensure all features operate as intended. Test cases will cover:')
    func_tests = [
        "User authentication and role-based access control",
        "CRUD operations for all content types",
        "Map functionality (markers, popups, filters)",
        "Search and filtering accuracy",
        "Image upload and display",
        "Content approval workflow",
        "Form import/export functionality"
    ]
    for test in func_tests:
        doc.add_paragraph(test, style='List Bullet')

    doc.add_heading('Usability Testing', level=3)
    add_para('Usability will be evaluated using standardized metrics and user feedback:')
    usability_tests = [
        "System Usability Scale (SUS) questionnaire",
        "Task completion rate and time-on-task measurements",
        "Error rate during common operations",
        "User satisfaction surveys (5-point Likert scale)",
        "Accessibility compliance (WCAG 2.1 guidelines)"
    ]
    for test in usability_tests:
        doc.add_paragraph(test, style='List Bullet')

    doc.add_heading('Performance Testing', level=3)
    add_para('Performance metrics will be measured to ensure acceptable response times:')
    perf_tests = [
        "Page load times (target: <3 seconds)",
        "Database query response times",
        "Concurrent user handling capacity",
        "Image optimization and loading performance",
        "API endpoint response times"
    ]
    for test in perf_tests:
        doc.add_paragraph(test, style='List Bullet')

    doc.add_heading('Security Testing', level=3)
    add_para('Security vulnerabilities will be identified and addressed:')
    security_tests = [
        "SQL injection prevention",
        "Cross-site scripting (XSS) protection",
        "Cross-site request forgery (CSRF) tokens",
        "Password hashing and encryption (Werkzeug/BCrypt)",
        "Session management and timeout",
        "Input validation and sanitization"
    ]
    for test in security_tests:
        doc.add_paragraph(test, style='List Bullet')

    doc.add_heading('Evaluation Plan', level=3)
    add_para('The system evaluation will follow a structured approach:')
    eval_plan = [
        ("Phase 1: Alpha Testing", "Internal testing by developers and LGU IT staff to identify critical bugs"),
        ("Phase 2: Beta Testing", "Limited release to selected Barangay Representatives for real-world feedback"),
        ("Phase 3: User Acceptance Testing", "Formal evaluation by Tourism Office staff using standardized questionnaires"),
        ("Phase 4: Performance Benchmarking", "Load testing and optimization based on metrics"),
        ("Phase 5: Final Deployment", "Production launch with monitoring and support")
    ]
    for phase, desc in eval_plan:
        p = doc.add_paragraph(style='List Number')
        p.add_run(phase).bold = True
        p.add_run(f" – {desc}")

    # ==================== APPENDICES ====================
    doc.add_page_break()
    doc.add_heading('Appendices', level=1)

    doc.add_heading('Appendix A: Administrative Guide', level=2)
    add_para('This guide is for the LGU Tourism Office and system administrators responsible for maintaining the platform.')

    doc.add_heading('Admin Dashboard', level=3)
    add_para('Access the dashboard at /admin/dashboard. The admin dashboard provides a high-level overview of system metrics, pending reviews, and recent activities.')

    doc.add_heading('Key Administrative Functions', level=3)
    admin_funcs = [
        ("Content Moderation", "Review and approve/reject submissions from Barangay Representatives. Access via /admin/content-review."),
        ("User Management", "Create, edit, or deactivate user accounts. Assign roles and manage permissions. Access via /admin/users."),
        ("Heritage Data", "Manage the five heritage categories: Natural Resources, Built Heritage, Movable Heritage, Intangible Heritage, and Living Heritage. Access via /admin/heritage."),
        ("System Settings", "Configure categories, tags, and global system parameters. Access via /admin/settings."),
        ("Reports", "Generate reports on system usage, content statistics, and user activity. Access via /admin/reports.")
    ]
    for func, desc in admin_funcs:
        p = doc.add_paragraph(style='List Number')
        p.add_run(func).bold = True
        p.add_run(f" – {desc}")

    doc.add_heading('Appendix B: User Manual', level=2)
    add_para('Welcome to the Interactive Digital Cultural Map of Mangatarem! This guide helps you navigate and explore the town\'s rich cultural and tourism assets.')

    doc.add_heading('Exploring the Map', level=3)
    add_para('The interactive map is the heart of the system.')
    map_guide = [
        "Access the map by clicking 'Map' in the main navigation",
        "Use mouse or touch to pan and zoom",
        "Click on colored markers to view attraction details",
        "Use the filter panel to show/hide categories",
        "Search for specific locations using the search bar",
        "Click 'Get Directions' for navigation assistance"
    ]
    for step in map_guide:
        doc.add_paragraph(step, style='List Bullet')

    doc.add_heading('Viewing Attractions', level=3)
    add_para('Each attraction has a dedicated detail page with:')
    attraction_info = [
        "High-quality photos in a gallery format",
        "Detailed description and historical context",
        "Location coordinates and address",
        "Opening hours and contact information",
        "Related attractions and suggested routes"
    ]
    for info in attraction_info:
        doc.add_paragraph(info, style='List Bullet')

    doc.add_heading('Appendix C: API Reference', level=2)
    add_para('This document provides a reference for the public and administrative API endpoints available in the system.')

    doc.add_heading('Authentication', level=3)
    add_para('Administrative and contributor endpoints require authentication. Use the session cookie provided after a successful login at /auth/login.')

    doc.add_heading('Public Endpoints', level=3)
    public_endpoints = [
        "GET /api/attractions - List all attractions",
        "GET /api/attractions/<id> - Get attraction details",
        "GET /api/events - List all events",
        "GET /api/heritage - Search heritage database",
        "GET /api/map/data - Get map marker data"
    ]
    for endpoint in public_endpoints:
        doc.add_paragraph(endpoint, style='List Bullet')

    doc.add_heading('Administrative Endpoints', level=3)
    admin_endpoints = [
        "POST /api/admin/attractions - Create attraction",
        "PUT /api/admin/attractions/<id> - Update attraction",
        "DELETE /api/admin/attractions/<id> - Delete attraction",
        "POST /api/admin/review - Approve/reject submission",
        "GET /api/admin/users - List all users"
    ]
    for endpoint in admin_endpoints:
        doc.add_paragraph(endpoint, style='List Bullet')

    doc.add_heading('Appendix D: Deployment Guide', level=2)
    add_para('This guide provides the steps to deploy and maintain the Interactive Digital Cultural Map in a production environment.')

    doc.add_heading('Prerequisites', level=3)
    prereqs = [
        "Supabase Account: For the PostgreSQL database",
        "Vercel Account: For serverless hosting",
        "Git: For version control and deployment",
        "Python 3.9+: For local development"
    ]
    for prereq in prereqs:
        doc.add_paragraph(prereq, style='List Bullet')

    doc.add_heading('Deployment Steps', level=3)
    deploy_steps = [
        "Set up Supabase project and run database migrations",
        "Configure environment variables (DATABASE_URL, SECRET_KEY)",
        "Connect Vercel to GitHub repository",
        "Deploy frontend to Vercel",
        "Configure custom domain (optional)",
        "Set up monitoring and logging",
        "Perform post-deployment testing"
    ]
    for i, step in enumerate(deploy_steps, 1):
        doc.add_paragraph(f"{i}. {step}", style='List Number')

    doc.add_heading('Appendix E: Core System Overview', level=2)
    add_para('The Interactive Digital Cultural Map and Local Tourism Information System for Mangatarem, Pangasinan is a GIS-based platform designed to document, map, and promote the cultural, historical, and tourism assets of Mangatarem.')

    doc.add_heading('System Goals', level=3)
    system_goals = [
        ("Documentation", "Digitally preserve Mangatarem's cultural heritage"),
        ("Accessibility", "Provide easy access to tourism information"),
        ("Promotion", "Enhance visibility of local attractions"),
        ("Education", "Support academic research and learning"),
        ("Preservation", "Document traditions for future generations")
    ]
    for goal, desc in system_goals:
        p = doc.add_paragraph(style='List Number')
        p.add_run(goal).bold = True
        p.add_run(f" – {desc}")

    doc.add_heading('Target Users', level=3)
    users = [
        "Tourists and visitors exploring Mangatarem",
        "LGU Tourism Office staff managing content",
        "Barangay Representatives contributing data",
        "Students and researchers studying local culture",
        "Residents interested in local heritage"
    ]
    for user in users:
        doc.add_paragraph(user, style='List Bullet')

    # Save the document
    output_path = r'D:\porjects\Interactive-Digital-Cultural-Map-and-Local-Tourism-Information-system-for-Mangatarem--Pangasinan\docs\PROJECT_DOCUMENTATION.docx'
    doc.save(output_path)
    print(f"Document saved successfully to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_document()
