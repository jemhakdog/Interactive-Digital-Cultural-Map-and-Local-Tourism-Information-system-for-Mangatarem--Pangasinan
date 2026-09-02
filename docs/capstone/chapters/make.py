import os
import re

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor


def create_element(name):
    return OxmlElement(name)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '<w:bottom w:val="single" w:sz="6" w:space="0" w:color="333333"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        '<w:insideV w:val="none"/>'
        '<w:left w:val="none"/>'
        '<w:right w:val="none"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

def parse_runs(paragraph_obj, markdown_text):
    """Parses bold formatting **text** and adds runs to paragraph_obj."""
    parts = re.split(r'(\*\*.*?\*\*)', markdown_text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            bold_text = part[2:-2]
            run = paragraph_obj.add_run(bold_text)
            run.bold = True
        else:
            if part:
                paragraph_obj.add_run(part)

def compile_manuscript():
    doc = Document()
    
    # Configure 1-inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Set default Normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)

    # Configure Heading styles to be Black and Times New Roman
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = 'Times New Roman'
        heading_style.font.color.rgb = RGBColor(0, 0, 0)
        heading_style.font.bold = True
        if i == 1:
            heading_style.font.size = Pt(16)
        elif i == 2:
            heading_style.font.size = Pt(14)
        else:
            heading_style.font.size = Pt(12)

    chapters_dir = r"d:\porjects\capstone_system\docs\capstone\chapters"
    chapter_files = [
        "Chapter-1-Introduction.md",
        "Chapter-2-Methodology-and-Design.md",
        "Chapter-3-Results-and-Discussion.md"
    ]

    first_chapter = True

    for cfile in chapter_files:
        cpath = os.path.join(chapters_dir, cfile)
        if not os.path.exists(cpath):
            print(f"Warning: File {cpath} not found.")
            continue

        if not first_chapter:
            doc.add_page_break()
        first_chapter = False

        print(f"Compiling {cfile}...")
        with open(cpath, "r", encoding="utf-8") as f:
            lines = f.readlines()

        in_table = False
        table_rows = []
        table_headers = []
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            # Handle empty lines
            if not line:
                i += 1
                continue

            # Check if line is a table row
            if line.startswith('|') and line.endswith('|'):
                in_table = True
                row_cells = [cell.strip() for cell in line.split('|')[1:-1]]
                table_rows.append(row_cells)
                i += 1
                continue
            elif in_table:
                # We reached the end of the table
                # Process table rows
                if len(table_rows) > 0:
                    # Filter out separator row (contains ---)
                    valid_rows = []
                    for row in table_rows:
                        if all(re.match(r'^:?-+:?$', cell) or not cell for cell in row):
                            continue
                        valid_rows.append(row)
                    
                    if valid_rows:
                        num_cols = max(len(row) for row in valid_rows)
                        table = doc.add_table(rows=len(valid_rows), cols=num_cols)
                        table.autofit = True
                        set_table_borders(table)
                        
                        # Populate table
                        for row_idx, row_data in enumerate(valid_rows):
                            row = table.rows[row_idx]
                            
                            # Check if header row
                            is_header = (row_idx == 0)
                            
                            for col_idx, cell_data in enumerate(row_data):
                                if col_idx < len(row.cells):
                                    cell = row.cells[col_idx]
                                    cell.text = ""  # Clear default
                                    p = cell.paragraphs[0]
                                    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                                    
                                    # Formatted runs
                                    parse_runs(p, cell_data)
                                    
                                    # Styling
                                    set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
                                    if is_header:
                                        set_cell_shading(cell, "F2F2F2")
                                        for run in p.runs:
                                            run.bold = True
                                            run.font.size = Pt(11)
                                    else:
                                        for run in p.runs:
                                            run.font.size = Pt(10.5)
                
                # Reset table state
                in_table = False
                table_rows = []
                
            # Handle Headers
            if line.startswith('# '):
                heading_text = line[2:].strip()
                doc.add_heading(heading_text, level=1)
            elif line.startswith('## '):
                heading_text = line[3:].strip()
                doc.add_heading(heading_text, level=2)
            elif line.startswith('### '):
                heading_text = line[4:].strip()
                doc.add_heading(heading_text, level=3)
            # Handle Bullet Lists
            elif line.startswith(('- ', '* ')):
                list_text = line[2:].strip()
                p = doc.add_paragraph(style='List Bullet')
                parse_runs(p, list_text)
            # Handle Numbered Lists
            elif re.match(r'^\d+\.\s', line):
                match = re.match(r'^(\d+)\.\s(.*)', line)
                list_text = match.group(2).strip()
                p = doc.add_paragraph(style='List Number')
                parse_runs(p, list_text)
            # Handle Mermaid placeholders or other generic codeblocks (skip raw mermaid)
            elif line.startswith('```'):
                # Skip codeblocks
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    i += 1
            # Skip image placeholders
            elif line.startswith(('<put the image', '[Insert', '*(Figure')):
                p = doc.add_paragraph(style='Normal')
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                parse_runs(p, line)
                for run in p.runs:
                    run.italic = True
                    run.font.size = Pt(11)
            # Standard Paragraph
            else:
                p = doc.add_paragraph(style='Normal')
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                parse_runs(p, line)
            
            i += 1

    output_path = os.path.join(chapters_dir, "Chapter_1_to_3_Consolidated.docx")
    doc.save(output_path)
    print(f"Compilation complete. Consolidated Word file saved to: {output_path}")

if __name__ == "__main__":
    compile_manuscript()