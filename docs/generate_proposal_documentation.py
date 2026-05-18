from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH


def create_document():
    """Create comprehensive project proposal documentation in DOCX format."""
    doc = Document()

    # Set default font for the document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)

    # Configure Heading styles
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = 'Times New Roman'
        heading_style.font.color.rgb = RGBColor(0, 0, 0)
        heading_style.font.bold = True
        if i == 1:
            heading_style.font.size = Pt(16)
        elif i == 2:
            heading_style.font.size = Pt(14)
        else:
            heading_style.font.size = Pt(12)

    # Helper to add styled paragraphs
    def add_para(text, style='Normal', bold=False, 
                 alignment=WD_PARAGRAPH_ALIGNMENT.JUSTIFY):
        p = doc.add_paragraph(text, style=style)
        p.alignment = alignment
        if bold:
            for run in p.runs:
                run.bold = True
        return p

    # ==================== TITLE PAGE ====================
    # Add title page
    for _ in range(3):
        doc.add_paragraph()
    
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    title_run = title_para.add_run(
        "INTERACTIVE DIGITAL CULTURAL MAP AND LOCAL TOURISM "
        "INFORMATION SYSTEM FOR MANGATAREM, PANGASINAN"
    )
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = 'Times New Roman'
    
    # Add subtitle
    for _ in range(2):
        doc.add_paragraph()
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Project Proposal Document")
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.name = 'Times New Roman'
    
    # Add team members
    for _ in range(4):
        doc.add_paragraph()
    
    team_para = doc.add_paragraph()
    team_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    team_run = team_para.add_run("Team B3:\n")
    team_run.bold = True
    team_run.font.size = Pt(12)
    
    members = [
        "Jem Carlo Austria",
        "Maryjane Dalas",
        "Rea Solis",
        "Joy De Guzman"
    ]
    for member in members:
        member_para = doc.add_paragraph()
        member_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        member_para.add_run(member)
    
    # Add institution info
    for _ in range(2):
        doc.add_paragraph()
    
    inst_para = doc.add_paragraph()
    inst_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    inst_para.add_run("In partnership with:\n")
    inst_para.add_run("LGU Mangatarem Tourism Office").bold = True
    
    # Add contact info
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.add_run("Rachelle T. Cabornay\n")
    contact_para.add_run("Senior Tourism Operations Officer\n")
    contact_para.add_run("Municipal Economic Enterprise Office (M.E.O.)")
    
    # Add date
    for _ in range(3):
        doc.add_paragraph()
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run("March 2026")
    
    # Page break after title
    doc.add_page_break()

    # ==================== TABLE OF CONTENTS ====================
    toc_para = doc.add_heading("Table of Contents", level=1)
    toc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    toc_sections = [
        "1. Project Overview",
        "2. Target Audience and Stakeholders",
        "3. Issues and Challenges Identified",
        "4. Main Objective and Purpose",
        "5. Scope of the System",
        "6. Users of the System",
        "7. Key Use Cases",
        "8. Functional Requirements",
        "9. Technical Requirements",
        "10. Design System",
        "11. Methodology",
        "12. System Flow and Logic",
        "13. Significance of the Study",
        "14. Implementation Timeline"
    ]
    
    for section in toc_sections:
        p = doc.add_paragraph(section)
        p.runs[0].bold = True
    
    doc.add_page_break()

    # ==================== SECTION 1: PROJECT OVERVIEW ====================
    doc.add_heading('1. Project Overview', level=1)

    project_info = [
        ("Project Name:", "Interactive Digital Cultural Map and Local Tourism Information System"),
        ("Type:", "Web Application"),
        ("Purpose:", "This platform helps users find cultural spots and tourism highlights in the municipality. It consolidates data onto one interactive map to boost local pride, assist tourists, and aid student research."),
        ("Problem Statement:", "The LGU Mangatarem Tourism Office currently faces challenges with fragmented, manual information management and irregularly updated content. This leads to inconsistent data and difficulties in disseminating accurate tourism information, hindering the promotion of local heritage.")
    ]
    
    for label, value in project_info:
        p = doc.add_paragraph()
        p.add_run(label).bold = True
        p.add_run(f" {value}")

    # ==================== SECTION 2: TARGET AUDIENCE ====================
    doc.add_heading('2. Target Audience and Stakeholders', level=1)

    add_para('The system serves multiple user groups with distinct needs and responsibilities:')

    audience_table = doc.add_table(rows=1, cols=2)
    audience_table.style = 'Table Grid'
    
    # Header row
    header_row = audience_table.rows[0].cells
    header_row[0].text = 'Target Audience'
    header_row[1].text = 'Needs and Requirements'
    
    # Make header bold
    for cell in header_row:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Add rows
    audience_data = [
        ("Tourists", "Visitors needing navigation, food recommendations, and attraction details"),
        ("Residents", "Locals interested in community events or discovering other barangays"),
        ("Students/Researchers", "Users looking for historical data and cultural heritage info"),
        ("LGU Tourism Office (Admins)", "Staff responsible for approving content and monitoring system health"),
        ("Barangay Representatives", "Designated individuals from each of the 82 barangays responsible for submitting local updates")
    ]
    
    for role, needs in audience_data:
        row = audience_table.add_row()
        row.cells[0].text = role
        row.cells[1].text = needs

    doc.add_paragraph()

    # ==================== SECTION 3: ISSUES AND CHALLENGES ====================
    doc.add_heading('3. Issues and Challenges Identified', level=1)

    add_para('The current tourism information management process faces several critical challenges:')

    challenges = [
        ("Fragmented Information Management", "Information management is currently fragmented and largely manual, resulting in irregularly updated online content."),
        ("Lack of Standardization", "Lack of standardized tourism materials leads to inconsistent data across different platforms."),
        ("Slow Coordination", "Slow coordination with stakeholders relying on traditional communication methods delays the sharing of accurate information."),
        ("Limited Accessibility", "Limited accessibility for students and researchers seeking cultural information.")
    ]
    
    for title, desc in challenges:
        p = doc.add_paragraph(style='List Number')
        p.add_run(title).bold = True
        p.add_run(f" – {desc}")

    # Current Process Flow
    add_para('Current Inefficient Process:', bold=True, alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)
    current_process = doc.add_paragraph()
    current_process.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    current_process.add_run("Search Social Media → Visit Tourism Office → Search Physical Files → High Risk of Confusion")
    current_process.runs[0].italic = True

    doc.add_page_break()

    # ==================== SECTION 4: MAIN OBJECTIVE ====================
    doc.add_heading('4. Main Objective and Purpose', level=1)

    add_para('The primary objective of this capstone project is to design and develop the Interactive Digital Cultural Map and Local Tourism Information System, a centralized web-based platform dedicated to showcasing the cultural identity and tourism assets of Mangatarem, Pangasinan.')

    add_para('This system addresses the information gap by:')
    objective_points = [
        "Integrating an interactive mapping interface with a comprehensive information portal.",
        "Allowing users to easily locate and explore barangay-level cultural spots, historical landmarks, and local festivities."
    ]
    for point in objective_points:
        doc.add_paragraph(point, style='List Bullet')

    add_para('Built using cost-effective and scalable open-source technologies like Python (Flask) and Leaflet.js/Mapbox, the system provides a viable solution for implementation by the Local Government Unit.')

    # ==================== SECTION 5: SCOPE ====================
    doc.add_heading('5. Scope of the System', level=1)

    add_para('The system encompasses the following core features and functionalities:')

    scope_features = [
        ("Interactive Cultural Map", "Navigable map displaying barangay-level cultural highlights with zoom, pan, and location details."),
        ("Cultural Information Portal", "Central repository for barangay profiles, cultural assets, traditions, and points of interest."),
        ("Events and Festival Directory", "Calendar view of local festivities and citywide celebrations with scheduling details."),
        ("Multimedia Gallery", "Organized collection of photos and videos showcasing local traditions and community activities."),
        ("Search and Filter Tools", "Quick retrieval of information by barangay, category, or attraction type."),
        ("Admin and Contributor Module", "Content management with approval workflow for authorized barangay representatives."),
        ("Tourist Guide and Routes", "Recommended themed itineraries for exploring the municipality efficiently."),
        ("Analytics Dashboard", "Visualizes system usage data and engagement trends for decision-makers.")
    ]
    
    for feature, desc in scope_features:
        p = doc.add_paragraph(style='List Number')
        p.add_run(feature).bold = True
        p.add_run(f" – {desc}")

    add_para('Extended Scope Details:')
    extended_scope = [
        "Covers all 82 barangays across the municipality.",
        "Enables digital coordination for all 82 barangays to independently submit and update their local cultural events and profiles.",
        "Promotes eco-tourism initiatives and hosts festivals like the Tupig Festival.",
        "Facilitates organized planning and efficient sharing of specific cultural events directly from the source.",
        "Real-time fare budget estimation: Enables tourists to select attractions and services to instantly generate an itemized cost breakdown (e.g., tricycle/jeepney fares)."
    ]
    for item in extended_scope:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ==================== SECTION 6: USERS ====================
    doc.add_heading('6. Users of the System', level=1)

    users_table = doc.add_table(rows=1, cols=2)
    users_table.style = 'Table Grid'
    
    # Header
    header_row = users_table.rows[0].cells
    header_row[0].text = 'User Group'
    header_row[1].text = 'Roles and Responsibilities'
    
    for cell in header_row:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    users_data = [
        ("System Administrator (Tourism Office Staff / IT Staff)", 
         "• Manages user accounts and system access permissions\n• Approves/rejects content submissions from contributors\n• Oversees platform maintenance and technical operations"),
        ("Contributor (Barangay Representative)", 
         "• Submits content about local attractions and events\n• Updates information within their jurisdiction\n• Uploads photos and videos of community activities"),
        ("Public User (Tourists / Visitors)", 
         "• Navigates interactive map to locate attractions\n• Searches and filters for specific points of interest\n• Views suggested routes and cultural information"),
        ("Academic Users (Students & Researchers)", 
         "• Accesses historical data and cultural profiles\n• Gathers research data on local heritage\n• Studies cultural traditions and community practices"),
        ("Stakeholders (LGU of Mangatarem - Primary)", 
         "• Main beneficiary and decision-maker for tourism promotion and cultural data management")
    ]
    
    for user, roles in users_data:
        row = users_table.add_row()
        row.cells[0].text = user
        row.cells[1].text = roles

    doc.add_paragraph()

    # ==================== SECTION 7: KEY USE CASES ====================
    doc.add_heading('7. Key Use Cases', level=1)

    use_cases = [
        ("Tourists Exploring Local Attractions", 
         "Navigate the Interactive Map to locate specific spots (e.g., Manleluag Spring), view pop-up details, get directions, and discover landmarks."),
        ("Academic Research and Data Gathering", 
         "Students access barangay profiles for academic studies, historical data, traditions, and local practices."),
        ("Content Contribution by Barangay Officials", 
         "Authorized Representatives log in to update local information, upload photos, update history, and add events."),
        ("Content Moderation and System Management", 
         "System Administrator uses Admin Dashboard to review submissions, approve/reject content, and manage users.")
    ]
    
    for title, desc in use_cases:
        p = doc.add_paragraph(style='List Number')
        p.add_run(title).bold = True
        p.add_run(f" – {desc}")

    # ==================== SECTION 8: FUNCTIONAL REQUIREMENTS ====================
    doc.add_heading('8. Functional Requirements', level=1)

    doc.add_heading('8.1 Interactive Cultural Map', level=2)
    map_reqs = [
        "Map Interface: Base map rendered using Leaflet.js centered on Mangatarem.",
        "Pins and Markers: Color-coded icons for categories (Nature, Historical, Religious, Food).",
        "Pop-up Cards: Clicking a pin opens a card with the attraction name, description, and link to details.",
        "Navigation: Zoom/Pan controls; 'Fly To' animation when selecting a spot from the sidebar."
    ]
    for req in map_reqs:
        doc.add_paragraph(req, style='List Bullet')

    doc.add_heading('8.2 Cultural and Tourism Information Portal', level=2)
    portal_reqs = [
        "Barangay Profiles: Dedicated templates displaying history, cultural assets, and unique features per barangay.",
        "Details Page: Full page view for attractions including Description, Location, and Image.",
        "Categories: Database categorization for 'Eateries,' 'Landmarks,' 'Religious Sites,' etc."
    ]
    for req in portal_reqs:
        doc.add_paragraph(req, style='List Bullet')

    doc.add_heading('8.3 Events and Festival Directory', level=2)
    events_reqs = [
        "Event List: Chronological display of upcoming events.",
        "Event Details: Title, Date, Location, and Description.",
        "Status Indicators: Logic to display event status (e.g., Pending Approval vs. Approved)."
    ]
    for req in events_reqs:
        doc.add_paragraph(req, style='List Bullet')

    doc.add_heading('8.4 Multimedia Gallery', level=2)
    gallery_reqs = [
        "Media Grid: Responsive grid layout for photos and videos.",
        "Uploads: Contributors can upload images/videos which are stored in the static file system.",
        "Filtering: Filter media by Barangay or Media Type (Photo/Video)."
    ]
    for req in gallery_reqs:
        doc.add_paragraph(req, style='List Bullet')

    doc.add_heading('8.5 Search and Filter Tools', level=2)
    search_reqs = [
        "Search Bar: Text input to filter map results by name or keyword.",
        "Filters: Buttons/Dropdowns to filter by Category (Nature, Heritage) or specific Barangay."
    ]
    for req in search_reqs:
        doc.add_paragraph(req, style='List Bullet')

    doc.add_heading('8.6 Admin and Contributor Module', level=2)
    admin_reqs = [
        "Role-Based Access Control (RBAC): Admin has full access to approve/reject/delete all content; Contributor can create/edit their own content with submissions requiring Admin approval.",
        "Dashboard: Specialized views for Admins (system-wide stats) and Contributors (personal submissions).",
        "Authentication: Secure login and registration flows."
    ]
    for req in admin_reqs:
        doc.add_paragraph(req, style='List Bullet')

    doc.add_heading('8.7 Tourist Guide and Suggested Routes', level=2)
    route_reqs = [
        "Static Routes: Informational pages suggesting itineraries (e.g., 'Heritage Walk').",
        "Future Scope: Interactive line drawing on the map."
    ]
    for req in route_reqs:
        doc.add_paragraph(req, style='List Bullet')

    doc.add_heading('8.8 Analytics Dashboard', level=2)
    analytics_reqs = [
        "Content Stats: Counters for Total Attractions, Events, and Gallery items.",
        "Engagement Metrics: Insights on most-viewed locations to help identify popular spots and improve promotional strategies."
    ]
    for req in analytics_reqs:
        doc.add_paragraph(req, style='List Bullet')

    doc.add_page_break()

    # ==================== SECTION 9: TECHNICAL REQUIREMENTS ====================
    doc.add_heading('9. Technical Requirements', level=1)

    doc.add_heading('9.1 Development Requirements', level=2)
    dev_reqs = [
        "Hardware: Intel Core i5/i7, 8GB+ RAM",
        "Languages: Python, HTML5, CSS3, JavaScript",
        "Frameworks: Flask, Tailwind CSS, Leaflet.js or Mapbox",
        "Database: SQLite (development) and Supabase (PostgreSQL)",
        "Version Control: Git and GitHub",
        "Code Editor: VS Code"
    ]
    for req in dev_reqs:
        doc.add_paragraph(req, style='List Bullet')

    doc.add_heading('9.2 Deployment and User Access Requirements', level=2)
    deploy_reqs = [
        "Web Hosting Service: Vercel",
        "Database Server: SQLite or PostgreSQL for production",
        "Map Tile Service: OpenStreetMap or CartoDB via Leaflet or Mapbox",
        "Client Device: Smartphone, Tablet, Laptop, or Desktop Computer",
        "Web Browser: Chrome, Firefox, Safari, or Edge",
        "Connectivity: Stable Internet Connection or Mobile Data Service"
    ]
    for req in deploy_reqs:
        doc.add_paragraph(req, style='List Bullet')

    doc.add_heading('9.3 Database Models', level=2)
    add_para('The system utilizes 19+ core database models organized into four categories:')

    doc.add_heading('User Management', level=3)
    user_models = [
        "USER: Handles authentication and roles (admin, contributor, user)",
        "PASSWORD_RESET_TOKEN: Secure account recovery"
    ]
    for model in user_models:
        doc.add_paragraph(model, style='List Bullet')

    doc.add_heading('Tourism and Map Data', level=3)
    tourism_models = [
        "ATTRACTION: Geo-located points of interest (coordinates: latitude, longitude)",
        "EVENT: Community activities and festivals",
        "BARANGAY_INFO: Cultural background profiles for each barangay"
    ]
    for model in tourism_models:
        doc.add_paragraph(model, style='List Bullet')

    doc.add_heading('Cultural Heritage Registry (Forms 01A-07)', level=3)
    heritage_models = [
        "HERITAGE_PROFILE: Base registry model",
        "Form 01A-07 Details: 7 specialized detail models for Natural, Built, Movable, Intangible Heritage, Personality Profiles, Cultural Institutions, and LGU Programs"
    ]
    for model in heritage_models:
        doc.add_paragraph(model, style='List Bullet')

    doc.add_heading('User Engagement', level=3)
    engagement_models = [
        "FAVORITE: Personal bookmarks",
        "EVENT_INTEREST: Event RSVPs",
        "REVIEW: Community feedback and ratings",
        "ANALYTICS_PAGE_VIEW: Internal engagement telemetry"
    ]
    for model in engagement_models:
        doc.add_paragraph(model, style='List Bullet')

    # ==================== SECTION 10: DESIGN SYSTEM ====================
    doc.add_heading('10. Design System', level=1)

    add_para('The visual identity of GoMangatarem is rooted in the lush greenery of Mangatarem\'s landscapes and the warmth of its cultural heritage.')

    doc.add_heading('10.1 Color Palette', level=2)
    add_para('Primary (Nature & Heritage):')
    primary_colors = [
        "Forest Green (#14532D): Deep, grounding",
        "Lush Accents (#22C55E): Vibrant, growth",
        "Soil Brown (#451A03): Earth, roots"
    ]
    for color in primary_colors:
        doc.add_paragraph(color, style='List Bullet')

    add_para('Secondary (Atmosphere):')
    secondary_colors = [
        "Sky Tint (#F0F9FF): Light, airy",
        "Cloud Gray (#F9FAFB): Neutral background",
        "Deep Slate (#111827): Text & high contrast"
    ]
    for color in secondary_colors:
        doc.add_paragraph(color, style='List Bullet')

    add_para('Surfaces:')
    surface_colors = [
        "Glass: rgba(255, 255, 255, 0.7) with backdrop-blur: 12px",
        "Paper: #FFFFFF (Solid components)"
    ]
    for color in surface_colors:
        doc.add_paragraph(color, style='List Bullet')

    doc.add_heading('10.2 Typography', level=2)
    typography = [
        "Headings: Outfit or Inter (Sans-serif, Bold, Tracking: -0.025em)",
        "Body: Inter (Sans-serif, Regular/Light, Leading: 1.625)"
    ]
    for typo in typography:
        doc.add_paragraph(typo, style='List Bullet')

    doc.add_heading('10.3 UI Principles', level=2)
    ui_principles = [
        ("Glassmorphism", "Use backdrop blur for floating UI elements like filters, navigation, and overlays with backdrop-blur-md, bg-white/70, and border border-white/20."),
        ("Micro-animations", "Subtle fade-in with slide-up for grid items; Scale 105% hover for images; 300ms button transitions; Gentle scroll-linked parallax movement."),
        ("Masonry Layout", "Embrace organic flow of content with variable heights and natural spacing for media galleries.")
    ]
    for title, desc in ui_principles:
        p = doc.add_paragraph(style='List Number')
        p.add_run(title).bold = True
        p.add_run(f" – {desc}")

    doc.add_page_break()

    # ==================== SECTION 11: METHODOLOGY ====================
    doc.add_heading('11. Methodology: Rapid Application Development (RAD)', level=1)

    add_para('The project employs the Rapid Application Development (RAD) methodology, an agile-based approach characterized by its flexible, user-centric design.')

    doc.add_heading('Key Benefits of RAD:', level=2)
    rad_benefits = [
        "Rapid Prototyping: Quick creation of working models for stakeholder feedback",
        "Iterative Feedback: Continuous improvement based on user input",
        "User-Centered Design: Active stakeholder participation throughout development",
        "Reduced Development Time: Faster delivery through parallel development cycles",
        "Flexible Requirements: Adaptability to changing needs during development"
    ]
    for benefit in rad_benefits:
        doc.add_paragraph(benefit, style='List Bullet')

    doc.add_heading('RAD Phases:', level=2)
    rad_phases = [
        ("Requirements Planning", "Collaboration with LGU staff and users to identify core problems and define objectives, user roles, and system scope."),
        ("User Design (Prototyping)", "Creation of interactive prototypes and wireframes using Figma, presented to stakeholders for iterative feedback."),
        ("Construction", "Actual coding using HTML, CSS, JavaScript for frontend and Python/Flask for backend, building functional modules in iterative cycles."),
        ("Cutover", "Comprehensive testing, deployment to production environment, user training, and ongoing maintenance.")
    ]
    for phase, desc in rad_phases:
        p = doc.add_paragraph(style='List Number')
        p.add_run(phase).bold = True
        p.add_run(f" – {desc}")

    # ==================== SECTION 12: SYSTEM FLOW ====================
    doc.add_heading('12. System Flow and Logic', level=1)

    doc.add_heading('12.1 Legacy Process: Key Bottlenecks', level=2)
    bottlenecks = [
        ("Manual Surveys", "Staff use paper forms for on-site data collection."),
        ("Manual Entry", "Typing data into Word or Excel is slow and leads to errors."),
        ("Traditional Delivery", "Data is submitted via physical handovers or simple dashboards."),
        ("Slow Verification", "Manual approvals cause long delays if corrections are needed."),
        ("Isolated Storage", "Data stays in local databases, making public access difficult.")
    ]
    for title, desc in bottlenecks:
        p = doc.add_paragraph(style='List Number')
        p.add_run(title).bold = True
        p.add_run(f" – {desc}")

    doc.add_heading('12.2 Legacy Tourism Issues', level=2)
    tourism_issues = [
        ("Office Visits", "Tourists must physically visit the tourism office for information."),
        ("Broken Information", "Brochures and verbal directions are hard to follow and retain."),
        ("Manual Maps", "Paper maps make navigation difficult and error-prone."),
        ("No Updates", "Tourists only discover closures or changes upon arrival.")
    ]
    for title, desc in tourism_issues:
        p = doc.add_paragraph(style='List Number')
        p.add_run(title).bold = True
        p.add_run(f" – {desc}")

    doc.add_heading('12.3 Core System Flows', level=2)

    doc.add_heading('Login and Security Flows', level=3)
    login_flows = [
        "Standard Login: Checks roles to open the correct dashboard.",
        "Google Login: Fast one-tap access for visitors; automatically creates profiles.",
        "Password Reset: Secure email-based recovery links.",
        "Logging Out: Ends session, clears tokens, redirects to home page."
    ]
    for flow in login_flows:
        doc.add_paragraph(flow, style='List Bullet')

    doc.add_heading('Registration and Validation Logic', level=3)
    reg_flows = [
        "Identity Check: Prevents duplicate accounts by verifying username and email uniqueness.",
        "User Roles: Fast access for visitors, extra verification for contributors.",
        "One Representative Limit: Only one registered representative per barangay.",
        "Admin Approval: Contributor accounts remain 'Pending' until confirmed by administrator."
    ]
    for flow in reg_flows:
        doc.add_paragraph(flow, style='List Bullet')

    doc.add_heading('Map Exploration Flow', level=3)
    map_flows = [
        "Auto Loading: System pulls verified attractions instantly upon page load.",
        "Interactive Markers: Rendered by Leaflet.js with color-coded categories.",
        "Barangay Filters: Sidebar sorting by specific barangay areas.",
        "Fast Previews: Popups with photos and quick information on marker click.",
        "Full Detail Pages: Complete attraction pages accessible via 'View Details' button."
    ]
    for flow in map_flows:
        doc.add_paragraph(flow, style='List Bullet')

    doc.add_page_break()

    # ==================== SECTION 13: SIGNIFICANCE ====================
    doc.add_heading('13. Significance of the Study', level=1)

    add_para('The Interactive Digital Cultural Map and Local Tourism Information System holds particular significance for the following beneficiaries:')

    significance = [
        ("Local Government Unit (LGU) of Mangatarem", 
         "As the main beneficiary and decision-maker, the LGU will benefit from a robust platform for tourism promotion and cultural data management, enabling them to verify and publish accurate information efficiently."),
        ("System Administrators (Tourism Office Staff / IT Staff)", 
         "They will benefit from an administrative dashboard that simplifies the management of user accounts, access permissions, and the approval/rejection of content submissions from contributors."),
        ("Barangay Representatives (Contributors)", 
         "They will benefit from a dedicated portal to submit and update local content, photos, and videos, empowering them to showcase the attractions and events within their respective jurisdictions."),
        ("Public Users (Tourists / Visitors)", 
         "They will benefit from an interactive map that helps them easily locate attractions, search and filter points of interest, and view suggested routes and cultural information for a better travel experience."),
        ("Students and Researchers", 
         "They will benefit from reliable access to historical data, cultural profiles, and community practices, facilitating their academic research and data gathering."),
        ("Residents of Mangatarem", 
         "They will gain cultural pride and benefit from the preservation of their heritage through a secure, digital platform that documents their traditions.")
    ]
    
    for i, (beneficiary, desc) in enumerate(significance, 1):
        p = doc.add_paragraph(style='List Number')
        p.add_run(beneficiary).bold = True
        p.add_run(f" – {desc}")

    # ==================== SECTION 14: IMPLEMENTATION ====================
    doc.add_heading('14. Implementation Timeline', level=1)

    add_para('The project implementation follows a phased approach to ensure systematic development and deployment:')

    timeline_table = doc.add_table(rows=1, cols=2)
    timeline_table.style = 'Table Grid'
    
    # Header
    header_row = timeline_table.rows[0].cells
    header_row[0].text = 'Phase'
    header_row[1].text = 'Activities and Deliverables'
    
    for cell in header_row:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    timeline_data = [
        ("Phase 1: Requirements Planning", 
         "• Stakeholder interviews with LGU Tourism Office\n• Documentation of current processes\n• Definition of system requirements\n• Project scope finalization"),
        ("Phase 2: System Design", 
         "• Database schema design (ERD)\n• Process flow modeling (DFD)\n• UI/UX prototyping in Figma\n• Design system documentation"),
        ("Phase 3: Development", 
         "• Backend development (Flask routes, models)\n• Frontend development (HTML/CSS/JavaScript)\n• Interactive map integration (Leaflet.js)\n• Authentication system implementation"),
        ("Phase 4: Testing", 
         "• Functional testing of all features\n• Usability testing with stakeholders\n• Performance optimization\n• Security vulnerability assessment"),
        ("Phase 5: Deployment", 
         "• Database migration to production\n• Application deployment to Vercel\n• User training for administrators\n• Documentation handover"),
        ("Phase 6: Maintenance", 
         "• Ongoing technical support\n• Regular content updates\n• Performance monitoring\n• Feature enhancements based on feedback")
    ]
    
    for phase, activities in timeline_data:
        row = timeline_table.add_row()
        row.cells[0].text = phase
        row.cells[1].text = activities

    # ==================== CONCLUSION ====================
    doc.add_page_break()
    doc.add_heading('Conclusion', level=1)

    add_para('The Interactive Digital Cultural Map and Local Tourism Information System represents a comprehensive solution to the challenges faced by the LGU of Mangatarem in managing and promoting tourism information. By leveraging modern web technologies and adopting a user-centered design approach, this system will:')

    conclusion_points = [
        "Centralize fragmented tourism data into a single, accessible platform",
        "Enable real-time updates and standardized information across all barangays",
        "Enhance the tourist experience through interactive mapping and search capabilities",
        "Support academic research and cultural preservation efforts",
        "Promote local economic development through increased tourism visibility",
        "Foster cultural pride and heritage preservation among residents"
    ]
    for point in conclusion_points:
        doc.add_paragraph(point, style='List Bullet')

    add_para('With the active participation of stakeholders and the systematic implementation approach outlined in this proposal, the project is positioned to deliver a sustainable, scalable solution that will serve the municipality of Mangatarem for years to come.')

    # Save the document
    output_path = r'D:\porjects\Interactive-Digital-Cultural-Map-and-Local-Tourism-Information-system-for-Mangatarem--Pangasinan\docs\PROJECT_PROPOSAL.docx'
    doc.save(output_path)
    print(f"Proposal document saved successfully to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_document()
