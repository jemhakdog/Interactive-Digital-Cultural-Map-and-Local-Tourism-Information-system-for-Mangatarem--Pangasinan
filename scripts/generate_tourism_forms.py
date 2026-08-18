"""
Generate Tourism Office Sample Forms Document
Creates a single DOCX file with all forms needed for the
Interactive Digital Cultural Map & Local Tourism Information System.

Usage:
    uv run python scripts/generate_tourism_forms.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path

# -- Constants --
MUNICIPALITY = "Mangatarem, Pangasinan"
SYSTEM_NAME = "Interactive Digital Cultural Map & Local Tourism Information System"
OUTPUT_DIR = Path(__file__).resolve().parent.parent / "docs"
OUTPUT_FILE = OUTPUT_DIR / "Tourism_Office_Sample_Forms.docx"

BARANGAYS = [
    "Abueg", "Baguinay", "Banaoang", "Bantog", "Binalay",
    "Bobon", "Brgy. I (Pob.)", "Brgy. II (Pob.)", "Buenlag",
    "Caaringayan", "Cabanbanan", "Cabaruan", "Caboloan",
    "Calaoagan", "Calaogan", "Calanutian", "Calaocan",
    "Camanggaan", "Canarvatan", "Canitoan", "Gueguesangen",
    "Inerangan", "Isla", "Lawak Langka", "Mabini", "Malicer",
    "Malico", "Manleluag", "Mapandan", "Nilombot",
    "Palca", "Pantar", "Paraoir", "Patiquin", "Petal",
    "Pogo", "Polo", "San Vicente", "Sumabnit",
    "Tabuyoc", "Vacante"
]

ATTRACTION_CATEGORIES = [
    "Nature", "Historical", "Religious", "Cultural",
    "Recreational", "Culinary / Food", "Other"
]

EVENT_CATEGORIES = ["Religious", "Civic", "Entertainment", "Cultural", "Sports"]

VEHICLE_TYPES = ["Tricycle", "Jeepney", "Van", "Bus", "Motorcycle (Habal-habal)"]

FOOD_CATEGORIES = ["Budget (Carenderia)", "Mid-Range (Restaurant)", "Local Delicacy / Specialty"]

ACCOMMODATION_TYPES = ["Budget (Inn/Homestay)", "Mid-Range (Hotel)", "Resort", "Pension House"]

MISC_COST_TYPES = ["Tour Guide", "Souvenir", "Activity Fee", "Parking", "Other"]


def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading.append(shd)


def style_header_row(row, bg_color="1B4F72", font_color=RGBColor(255, 255, 255)):
    """Style a table header row with background color and white bold text."""
    for cell in row.cells:
        set_cell_shading(cell, bg_color)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = font_color
                run.font.size = Pt(10)


def add_field_row(table, label, value="", col_widths=None):
    """Add a label-value row to a table."""
    row = table.add_row()
    row.cells[0].text = label
    row.cells[0].paragraphs[0].runs[0].font.bold = True if row.cells[0].paragraphs[0].runs else None
    row.cells[1].text = value

    # Make label bold
    for p in row.cells[0].paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(10)
    for p in row.cells[1].paragraphs:
        for r in p.runs:
            r.font.size = Pt(10)
    return row


def add_blank_lines(table, n=3):
    """Add blank entry rows to a table."""
    for _ in range(n):
        row = table.add_row()
        for cell in row.cells:
            cell.text = ""


def create_header(doc):
    """Create the document title page / header."""
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MUNICIPAL TOURISM OFFICE")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(27, 79, 114)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Municipality of {MUNICIPALITY}")
    run.font.size = Pt(12)
    run.font.bold = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TOURISM DATA COLLECTION FORMS")
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(27, 79, 114)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"For: {SYSTEM_NAME}")
    run.font.size = Pt(10)
    run.font.italic = True

    doc.add_paragraph()

    # Table of contents
    p = doc.add_paragraph()
    run = p.add_run("TABLE OF CONTENTS")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(27, 79, 114)

    forms = [
        "Form 1: Tourist Attraction / Site Registration Form",
        "Form 2: Event / Festival Registration Form",
        "Form 3: Barangay Cultural Profile Form",
        "Form 4: Gallery / Media Submission Form",
        "Form 5A: Attraction Entrance Fee Form",
        "Form 5B: Transport Fare Data Form",
        "Form 5C: Food & Dining Cost Form",
        "Form 5D: Accommodation / Lodging Cost Form",
        "Form 5E: Miscellaneous Cost Form (Tour Guide, Souvenirs, etc.)",
        "Form 6: Tourism Statistics & Visitor Data Form",
    ]
    for i, form in enumerate(forms, 1):
        p = doc.add_paragraph(form, style="List Number")
        for run in p.runs:
            run.font.size = Pt(10)

    doc.add_page_break()


def form_attraction(doc):
    """Form 1: Tourist Attraction / Site Registration Form."""
    doc.add_heading("Form 1: Tourist Attraction / Site Registration Form", level=1)

    p = doc.add_paragraph()
    run = p.add_run("Purpose: ")
    run.font.bold = True
    run.font.size = Pt(10)
    run = p.add_run("Register and document tourist attractions/sites in Mangatarem for the Digital Cultural Map system.")
    run.font.size = Pt(10)

    doc.add_paragraph()

    # Main info table
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header = table.rows[0]
    header.cells[0].merge(header.cells[1])
    header.cells[0].text = "ATTRACTION / SITE INFORMATION"
    style_header_row(header)

    fields = [
        ("Attraction Name:", ""),
        ("Category:", f"☐ {' ☐ '.join(ATTRACTION_CATEGORIES)}"),
        ("Barangay:", ""),
        ("Complete Address / Landmark:", ""),
        ("GPS Coordinates (if available):", "Latitude: ____________  Longitude: ____________"),
        ("Description:", "\n\n\n"),
        ("Operating Hours:", ""),
        ("Contact Person / Number:", ""),
        ("Entrance Fee:", "☐ Free   ☐ Fixed: ₱________   ☐ Range: ₱________ to ₱________"),
        ("Fee Notes:", "(e.g., Free for students, senior discount, weekend rates)\n"),
        ("Status:", "☐ Open to Public   ☐ Seasonal   ☐ By Appointment   ☐ Under Development"),
    ]
    for label, value in fields:
        add_field_row(table, label, value)

    doc.add_paragraph()

    # Photo attachment section
    p = doc.add_paragraph()
    run = p.add_run("PHOTO ATTACHMENTS")
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(27, 79, 114)

    photo_table = doc.add_table(rows=3, cols=3)
    photo_table.style = "Table Grid"
    photo_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i in range(3):
        for j in range(3):
            cell = photo_table.cell(i, j)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"Photo {i * 3 + j + 1}\n(Attach here)")
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(150, 150, 150)

    doc.add_paragraph()

    # Certification
    p = doc.add_paragraph()
    run = p.add_run("Submitted by: ________________________     Date: ________________")
    run.font.size = Pt(10)
    p = doc.add_paragraph()
    run = p.add_run("Verified by (Tourism Officer): ________________________     Date: ________________")
    run.font.size = Pt(10)

    doc.add_page_break()


def form_event(doc):
    """Form 2: Event / Festival Registration Form."""
    doc.add_heading("Form 2: Event / Festival Registration Form", level=1)

    p = doc.add_paragraph()
    run = p.add_run("Purpose: ")
    run.font.bold = True
    run.font.size = Pt(10)
    run = p.add_run("Document local events, festivals, and celebrations for the Events Calendar feature.")
    run.font.size = Pt(10)

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header = table.rows[0]
    header.cells[0].merge(header.cells[1])
    header.cells[0].text = "EVENT / FESTIVAL INFORMATION"
    style_header_row(header)

    fields = [
        ("Event / Festival Name:", ""),
        ("Category:", f"☐ {' ☐ '.join(EVENT_CATEGORIES)}"),
        ("Description:", "\n\n\n"),
        ("Date(s):", "Start: ____________  End: ____________"),
        ("Is this a recurring event?", "☐ Yes (Every year)   ☐ No (One-time)   Schedule: ____________"),
        ("Time:", "From: ____________  To: ____________"),
        ("Venue / Location:", ""),
        ("Barangay:", ""),
        ("Organizer / Contact Person:", ""),
        ("Contact Number:", ""),
        ("Estimated Attendance:", ""),
        ("Activities / Highlights:", "\n\n"),
        ("Entrance Fee:", "☐ Free   ☐ Paid: ₱________"),
    ]
    for label, value in fields:
        add_field_row(table, label, value)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Event Photo(s) Attached: ☐ Yes   ☐ No     Number of photos: ________")
    run.font.size = Pt(10)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Submitted by: ________________________     Date: ________________")
    run.font.size = Pt(10)
    p = doc.add_paragraph()
    run = p.add_run("Verified by (Tourism Officer): ________________________     Date: ________________")
    run.font.size = Pt(10)

    doc.add_page_break()


def form_barangay_profile(doc):
    """Form 3: Barangay Cultural Profile Form."""
    doc.add_heading("Form 3: Barangay Cultural Profile Form", level=1)

    p = doc.add_paragraph()
    run = p.add_run("Purpose: ")
    run.font.bold = True
    run.font.size = Pt(10)
    run = p.add_run("Collect cultural and historical information for each barangay's profile page in the system.")
    run.font.size = Pt(10)

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header = table.rows[0]
    header.cells[0].merge(header.cells[1])
    header.cells[0].text = "BARANGAY CULTURAL PROFILE"
    style_header_row(header)

    fields = [
        ("Barangay Name:", ""),
        ("Barangay Captain:", ""),
        ("Contact Number:", ""),
        ("Population (approx.):", ""),
        ("Brief History:", "\n\n\n\n\n"),
        ("Cultural Assets:", "(Landmarks, heritage buildings, monuments, etc.)\n\n\n"),
        ("Traditions & Festivals:", "(Annual celebrations, fiestas, rituals, etc.)\n\n\n"),
        ("Local Practices:", "(Farming, fishing, weaving, crafts, cuisine, etc.)\n\n\n"),
        ("Unique Features:", "(What makes this barangay special? Natural features, stories, etc.)\n\n\n"),
        ("Tourist Spots in this Barangay:", "\n\n"),
        ("Number of Tourist Spots:", ""),
    ]
    for label, value in fields:
        add_field_row(table, label, value)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Accomplished by: ________________________     Position: ________________________")
    run.font.size = Pt(10)
    p = doc.add_paragraph()
    run = p.add_run("Date: ________________")
    run.font.size = Pt(10)

    # List of barangays reference
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Reference: List of Barangays in Mangatarem")
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(27, 79, 114)

    brgy_table = doc.add_table(rows=1, cols=4)
    brgy_table.style = "Table Grid"
    brgy_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header = brgy_table.rows[0]
    for i, h in enumerate(["#", "Barangay Name", "#", "Barangay Name"]):
        header.cells[i].text = h
    style_header_row(header)

    half = (len(BARANGAYS) + 1) // 2
    for i in range(half):
        row = brgy_table.add_row()
        row.cells[0].text = str(i + 1)
        row.cells[1].text = BARANGAYS[i]
        if i + half < len(BARANGAYS):
            row.cells[2].text = str(i + half + 1)
            row.cells[3].text = BARANGAYS[i + half]
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_page_break()


def form_gallery(doc):
    """Form 4: Gallery / Media Submission Form."""
    doc.add_heading("Form 4: Gallery / Media Submission Form", level=1)

    p = doc.add_paragraph()
    run = p.add_run("Purpose: ")
    run.font.bold = True
    run.font.size = Pt(10)
    run = p.add_run("Track photos and videos submitted for the multimedia gallery feature.")
    run.font.size = Pt(10)

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header = table.rows[0]
    header.cells[0].merge(header.cells[1])
    header.cells[0].text = "MEDIA SUBMISSION DETAILS"
    style_header_row(header)

    fields = [
        ("Submitted by (Name):", ""),
        ("Contact Number / Email:", ""),
        ("Date of Submission:", ""),
        ("Media Type:", "☐ Photo   ☐ Video"),
        ("Number of Files:", ""),
        ("Caption / Description:", "\n\n"),
        ("Location / Subject:", ""),
        ("Barangay:", ""),
        ("Date Taken:", ""),
        ("Source / Credit:", "(Who took the photo/video?)\n"),
        ("Permission to Use:", "☐ Yes, I authorize the use of this media for the Digital Cultural Map system."),
    ]
    for label, value in fields:
        add_field_row(table, label, value)

    doc.add_paragraph()

    # Media log table
    p = doc.add_paragraph()
    run = p.add_run("MEDIA LOG (for multiple submissions)")
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(27, 79, 114)

    log_table = doc.add_table(rows=1, cols=5)
    log_table.style = "Table Grid"
    log_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["#", "File Name", "Type (Photo/Video)", "Caption", "Date Taken"]
    header_row = log_table.rows[0]
    for i, h in enumerate(headers):
        header_row.cells[i].text = h
    style_header_row(header_row)

    for i in range(8):
        row = log_table.add_row()
        row.cells[0].text = str(i + 1)
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Signature: ________________________     Date: ________________")
    run.font.size = Pt(10)

    doc.add_page_break()


def form_attraction_cost(doc):
    """Form 5A: Attraction Entrance Fee Form."""
    doc.add_heading("Form 5A: Attraction Entrance Fee Data Form", level=1)

    p = doc.add_paragraph()
    run = p.add_run("Purpose: ")
    run.font.bold = True
    run.font.size = Pt(10)
    run = p.add_run("Collect entrance fee data for each attraction for the Trip Cost Estimator feature.")
    run.font.size = Pt(10)

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["#", "Attraction Name", "Barangay", "Fee Type\n(Free/Fixed/Range)", "Fixed Fee\n(₱)", "Range Min\n(₱)", "Range Max\n(₱)"]
    header_row = table.rows[0]
    for i, h in enumerate(headers):
        header_row.cells[i].text = h
    style_header_row(header_row)

    for i in range(15):
        row = table.add_row()
        row.cells[0].text = str(i + 1)
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("Notes / Special Pricing Rules:")
    run.font.bold = True
    run.font.size = Pt(10)
    p = doc.add_paragraph("(e.g., Free for students, senior citizen discount, weekend rates, group rates)")
    for r in p.runs:
        r.font.size = Pt(9)
        r.font.italic = True

    notes_table = doc.add_table(rows=5, cols=1)
    notes_table.style = "Table Grid"

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Data provided by: ________________________     Date: ________________")
    run.font.size = Pt(10)

    doc.add_page_break()


def form_transport(doc):
    """Form 5B: Transport Fare Data Form."""
    doc.add_heading("Form 5B: Transport Fare Data Form", level=1)

    p = doc.add_paragraph()
    run = p.add_run("Purpose: ")
    run.font.bold = True
    run.font.size = Pt(10)
    run = p.add_run("Collect local transport fare estimates for the Trip Cost Estimator feature.")
    run.font.size = Pt(10)

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["#", "Vehicle Type", "Origin\n(From)", "Destination\n(To)", "Estimated Fare\n(₱)", "Notes"]
    header_row = table.rows[0]
    for i, h in enumerate(headers):
        header_row.cells[i].text = h
    style_header_row(header_row)

    for i in range(15):
        row = table.add_row()
        row.cells[0].text = str(i + 1)
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Vehicle Types: ")
    run.font.bold = True
    run.font.size = Pt(10)
    run = p.add_run(", ".join(VEHICLE_TYPES))
    run.font.size = Pt(10)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Common Routes (please fill fares for these):")
    run.font.bold = True
    run.font.size = Pt(10)

    routes = [
        "Town Center ↔ Manleluag Spring National Park",
        "Town Center ↔ Timmanguyob Falls",
        "Town Center ↔ ST. Raymund de Peñafort Church",
        "Town Center ↔ Public Market",
        "Poblacion ↔ Neighboring major barangays",
    ]
    for route in routes:
        p = doc.add_paragraph(route, style="List Bullet")
        for r in p.runs:
            r.font.size = Pt(9)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Data provided by: ________________________     Date: ________________")
    run.font.size = Pt(10)

    doc.add_page_break()


def form_food(doc):
    """Form 5C: Food & Dining Cost Form."""
    doc.add_heading("Form 5C: Food & Dining Cost Data Form", level=1)

    p = doc.add_paragraph()
    run = p.add_run("Purpose: ")
    run.font.bold = True
    run.font.size = Pt(10)
    run = p.add_run("Collect meal/dining price ranges for the Trip Cost Estimator feature.")
    run.font.size = Pt(10)

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = [
        "#", "Establishment Name", "Category", "Barangay",
        "Price Min\n(₱/meal)", "Price Max\n(₱/meal)", "Description / Specialty"
    ]
    header_row = table.rows[0]
    for i, h in enumerate(headers):
        header_row.cells[i].text = h
    style_header_row(header_row)

    for i in range(12):
        row = table.add_row()
        row.cells[0].text = str(i + 1)
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Categories: ")
    run.font.bold = True
    run.font.size = Pt(10)
    run = p.add_run(", ".join(FOOD_CATEGORIES))
    run.font.size = Pt(10)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Data provided by: ________________________     Date: ________________")
    run.font.size = Pt(10)

    doc.add_page_break()


def form_accommodation(doc):
    """Form 5D: Accommodation / Lodging Cost Form."""
    doc.add_heading("Form 5D: Accommodation / Lodging Cost Data Form", level=1)

    p = doc.add_paragraph()
    run = p.add_run("Purpose: ")
    run.font.bold = True
    run.font.size = Pt(10)
    run = p.add_run("Collect lodging/accommodation rates for the Trip Cost Estimator feature.")
    run.font.size = Pt(10)

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = [
        "#", "Establishment Name", "Type", "Barangay",
        "Rate / Night\n(₱)", "Contact Info", "Description"
    ]
    header_row = table.rows[0]
    for i, h in enumerate(headers):
        header_row.cells[i].text = h
    style_header_row(header_row)

    for i in range(10):
        row = table.add_row()
        row.cells[0].text = str(i + 1)
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Types: ")
    run.font.bold = True
    run.font.size = Pt(10)
    run = p.add_run(", ".join(ACCOMMODATION_TYPES))
    run.font.size = Pt(10)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Data provided by: ________________________     Date: ________________")
    run.font.size = Pt(10)

    doc.add_page_break()


def form_misc_cost(doc):
    """Form 5E: Miscellaneous Cost Form."""
    doc.add_heading("Form 5E: Miscellaneous Cost Data Form", level=1)

    p = doc.add_paragraph()
    run = p.add_run("Purpose: ")
    run.font.bold = True
    run.font.size = Pt(10)
    run = p.add_run("Collect tour guide fees, souvenir prices, and other miscellaneous costs.")
    run.font.size = Pt(10)

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = [
        "#", "Item / Service Name", "Cost Type", "Barangay",
        "Price Min\n(₱)", "Price Max\n(₱)", "Description"
    ]
    header_row = table.rows[0]
    for i, h in enumerate(headers):
        header_row.cells[i].text = h
    style_header_row(header_row)

    for i in range(12):
        row = table.add_row()
        row.cells[0].text = str(i + 1)
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Cost Types: ")
    run.font.bold = True
    run.font.size = Pt(10)
    run = p.add_run(", ".join(MISC_COST_TYPES))
    run.font.size = Pt(10)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Data provided by: ________________________     Date: ________________")
    run.font.size = Pt(10)

    doc.add_page_break()


def form_tourism_statistics(doc):
    """Form 6: Tourism Statistics & Visitor Data Form."""
    doc.add_heading("Form 6: Tourism Statistics & Visitor Data Form", level=1)

    p = doc.add_paragraph()
    run = p.add_run("Purpose: ")
    run.font.bold = True
    run.font.size = Pt(10)
    run = p.add_run("Collect tourism statistics to validate system analytics and support capstone findings.")
    run.font.size = Pt(10)

    doc.add_paragraph()

    # General stats
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header = table.rows[0]
    header.cells[0].merge(header.cells[1])
    header.cells[0].text = "GENERAL TOURISM STATISTICS"
    style_header_row(header)

    fields = [
        ("Year / Period Covered:", ""),
        ("Total No. of Tourists (Annual):", ""),
        ("Local Tourists:", ""),
        ("Foreign Tourists:", ""),
        ("Peak Season / Months:", ""),
        ("Most Visited Attraction:", ""),
        ("No. of Registered Tourist Spots:", ""),
        ("No. of Accredited Establishments:", ""),
        ("Tourism Revenue (if available):", ""),
    ]
    for label, value in fields:
        add_field_row(table, label, value)

    doc.add_paragraph()

    # Per-attraction visitor count
    p = doc.add_paragraph()
    run = p.add_run("VISITOR COUNT PER ATTRACTION (if available)")
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(27, 79, 114)

    vis_table = doc.add_table(rows=1, cols=4)
    vis_table.style = "Table Grid"
    vis_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["#", "Attraction Name", "Annual Visitors\n(Estimate)", "Notes"]
    header_row = vis_table.rows[0]
    for i, h in enumerate(headers):
        header_row.cells[i].text = h
    style_header_row(header_row)

    for i in range(10):
        row = vis_table.add_row()
        row.cells[0].text = str(i + 1)
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph()

    # Challenges
    p = doc.add_paragraph()
    run = p.add_run("CURRENT CHALLENGES IN TOURISM PROMOTION:")
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(27, 79, 114)

    chal_table = doc.add_table(rows=6, cols=1)
    chal_table.style = "Table Grid"
    chal_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    chal_table.rows[0].cells[0].text = "1."
    chal_table.rows[1].cells[0].text = "2."
    chal_table.rows[2].cells[0].text = "3."
    chal_table.rows[3].cells[0].text = "4."
    chal_table.rows[4].cells[0].text = "5."
    chal_table.rows[5].cells[0].text = "(Add more if needed)"

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Data provided by: ________________________     Position: ________________________")
    run.font.size = Pt(10)
    p = doc.add_paragraph()
    run = p.add_run("Date: ________________")
    run.font.size = Pt(10)


def set_document_styles(doc):
    """Configure global document styles."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Heading 1 style
    h1_style = doc.styles["Heading 1"]
    h1_style.font.size = Pt(14)
    h1_style.font.color.rgb = RGBColor(27, 79, 114)
    h1_style.font.bold = True

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)


def main():
    doc = Document()
    set_document_styles(doc)

    # Build the document
    create_header(doc)
    form_attraction(doc)
    form_event(doc)
    form_barangay_profile(doc)
    form_gallery(doc)
    form_attraction_cost(doc)
    form_transport(doc)
    form_food(doc)
    form_accommodation(doc)
    form_misc_cost(doc)
    form_tourism_statistics(doc)

    # Save
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT_FILE))
    print(f"[OK] Document generated: {OUTPUT_FILE}")
    print(f"     Contains 10 forms across {len(doc.element.body)} elements")


if __name__ == "__main__":
    main()
