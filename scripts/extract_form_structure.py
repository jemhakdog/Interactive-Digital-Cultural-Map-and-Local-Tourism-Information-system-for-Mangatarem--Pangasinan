"""
Extract structure from tourism office forms (DOCX) to identify database fields.
This script analyzes form templates to determine required database schema.
"""
from pathlib import Path
from docx import Document
import json


def extract_form_structure(docx_path):
    """Extract field structure from a DOCX form template."""
    doc = Document(docx_path)
    form_name = Path(docx_path).stem
    
    print(f"\n{'='*80}")
    print(f"ANALYZING: {form_name}")
    print(f"{'='*80}\n")
    
    structure = {
        "form_name": form_name,
        "paragraphs": [],
        "tables": [],
        "fields": []
    }
    
    # Extract paragraphs (headings, labels, etc.)
    print("PARAGRAPHS/HEADINGS:")
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            style = para.style.name
            structure["paragraphs"].append({
                "index": i,
                "text": text,
                "style": style
            })
            print(f"  [{style}] {text[:100]}")
    
    # Extract tables (form fields structure)
    print(f"\nTABLES ({len(doc.tables)} found):")
    for table_idx, table in enumerate(doc.tables):
        print(f"\n  Table {table_idx + 1} ({len(table.rows)} rows × {len(table.columns)} cols):")
        
        table_data = {
            "table_index": table_idx,
            "rows": len(table.rows),
            "columns": len(table.columns),
            "content": []
        }
        
        for row_idx, row in enumerate(table.rows):
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            table_data["content"].append(row_data)
            
            # Print first 5 rows to see structure
            if row_idx < 5:
                print(f"    Row {row_idx + 1}: {' | '.join(row_data[:3])}")
        
        if len(table.rows) > 5:
            print(f"    ... ({len(table.rows) - 5} more rows)")
        
        structure["tables"].append(table_data)
    
    return structure


def main():
    """Process all forms in the gathered_froms directory."""
    forms_dir = Path(__file__).parent.parent / "docs" / "interview_data" / "gathered_froms"
    
    if not forms_dir.exists():
        print(f"Directory not found: {forms_dir}")
        return
    
    # Get all DOCX files (excluding temp files starting with ~$)
    form_files = [f for f in forms_dir.glob("*.docx") if not f.name.startswith("~$")]
    
    print(f"\nFound {len(form_files)} forms to analyze\n")
    
    all_structures = {}
    
    for form_file in sorted(form_files):
        structure = extract_form_structure(form_file)
        all_structures[structure["form_name"]] = structure
    
    # Save to JSON for further analysis
    output_file = Path(__file__).parent.parent / "docs" / "interview_data" / "forms_structure_analysis.json"
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(all_structures, f, indent=2, ensure_ascii=False)
    
    print("\n" + "="*80)
    print("Analysis complete! Structure saved to:")
    print(f"   {output_file}")
    print("="*80 + "\n")


if __name__ == "__main__":
    main()
